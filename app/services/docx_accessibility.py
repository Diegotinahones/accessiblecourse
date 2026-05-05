from __future__ import annotations

import re
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from xml.etree import ElementTree
from zipfile import BadZipFile, ZipFile

from docx import Document

from app.core.config import Settings
from app.services.html_accessibility import (
    AccessibilityCheckResult,
    AccessibilityReport,
    append_accessibility_resource_result,
    load_accessibility_report,
    recompute_accessibility_summary,
    remove_accessibility_results,
    save_accessibility_report,
)
from app.services.resource_core import ResourceContentResult, get_resource_content, normalize_resource

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
GENERIC_DOCX_TITLES = {
    "document",
    "documento",
    "untitled",
    "sin titulo",
    "sin título",
    "word",
}
GENERIC_LINK_TEXTS = {
    "aqui",
    "aquí",
    "click aqui",
    "click aquí",
    "clic aqui",
    "clic aquí",
    "leer mas",
    "leer más",
    "mas",
    "más",
}
URL_TEXT_RE = re.compile(r"^(?:https?://|www\.)\S+$", re.IGNORECASE)
MANUAL_LIST_RE = re.compile(r"^\s*(?:[-*•]\s+|\d+[\.)]\s+)")


def analyze_docx_accessibility(resource: Any, docx_path: Path) -> list[AccessibilityCheckResult]:
    try:
        document = Document(str(docx_path))
        xml_context = _DocxXmlContext.from_path(docx_path)
    except (BadZipFile, OSError, ValueError, Exception) as exc:
        return [
            _result(
                "docx.open",
                "Documento legible",
                "FAIL",
                f"No se pudo abrir o parsear el DOCX: {exc.__class__.__name__}.",
                "Comprueba que el fichero sea un DOCX válido y no esté corrupto.",
            )
        ]

    context = _DocxContext.from_document(resource, docx_path, document, xml_context)
    return [
        _check_readable(),
        _check_extractable_text(context),
        _check_language(context),
        _check_title(context),
        _check_heading_styles(context),
        _check_heading_hierarchy(context),
        _check_image_alt(context),
        _check_tables(context),
        _check_links(context),
        _check_lists(context),
    ]


def run_docx_accessibility_scan(
    *,
    settings: Settings,
    job_id: str,
    resources: list[dict[str, Any]],
    canvas_client: Any | None = None,
    canvas_credentials: Any | None = None,
    course_id: str | None = None,
) -> AccessibilityReport:
    report = load_accessibility_report(settings, job_id)
    report.generatedAt = datetime.now(UTC)
    remove_accessibility_results(report, "DOCX")

    for resource in resources:
        core = normalize_resource(resource)
        if not _should_attempt_docx_scan(resource, core):
            continue

        try:
            content = get_resource_content(
                job_id,
                core.id,
                settings=settings,
                resources=resources,
                canvas_client=canvas_client,
                canvas_credentials=canvas_credentials,
                course_id=course_id,
            )
        except Exception as exc:
            checks = [
                _result(
                    "docx.content.available",
                    "Contenido DOCX disponible",
                    "ERROR",
                    f"No se pudo recuperar el DOCX: {exc.__class__.__name__}.",
                    "Comprueba que el documento Word sea accesible para el backend.",
                )
            ]
            append_accessibility_resource_result(report, resource, checks, analysis_type="DOCX")
            continue

        if not content.ok:
            checks = [
                _result(
                    "docx.content.available",
                    "Contenido DOCX disponible",
                    "ERROR",
                    content.errorDetail or "No se pudo recuperar el DOCX del recurso.",
                    "Comprueba que el documento Word sea accesible para el backend.",
                )
            ]
            append_accessibility_resource_result(report, resource, checks, analysis_type="DOCX")
            continue

        if not _content_is_docx(content):
            continue

        if not content.binaryPath:
            checks = [
                _result(
                    "docx.content.available",
                    "Contenido DOCX disponible",
                    "ERROR",
                    "El recurso DOCX no tiene ruta binaria segura para analizar.",
                    "Asegura que el DOCX se pueda resolver o cachear desde el backend.",
                )
            ]
            append_accessibility_resource_result(report, resource, checks, analysis_type="DOCX")
            continue

        checks = analyze_docx_accessibility(resource, Path(content.binaryPath))
        append_accessibility_resource_result(report, resource, checks, analysis_type="DOCX")

    recompute_accessibility_summary(report)
    save_accessibility_report(settings, job_id, report)
    return report


def ensure_docx_accessibility_report(
    *,
    settings: Settings,
    job_id: str,
    resources: list[dict[str, Any]],
    canvas_client: Any | None = None,
    canvas_credentials: Any | None = None,
    course_id: str | None = None,
) -> AccessibilityReport:
    report = load_accessibility_report(settings, job_id)
    eligible_count = sum(
        1 for resource in resources if _should_attempt_docx_scan(resource, normalize_resource(resource))
    )
    if eligible_count == 0 or report.summary.docxResourcesTotal >= eligible_count:
        return report
    return run_docx_accessibility_scan(
        settings=settings,
        job_id=job_id,
        resources=resources,
        canvas_client=canvas_client,
        canvas_credentials=canvas_credentials,
        course_id=course_id,
    )


class _DocxXmlContext:
    def __init__(
        self,
        *,
        languages: list[str],
        image_alt_values: list[str | None],
        links: list[str],
        table_header_count: int,
        structured_list_count: int,
        manual_list_count: int,
    ) -> None:
        self.languages = languages
        self.image_alt_values = image_alt_values
        self.links = links
        self.table_header_count = table_header_count
        self.structured_list_count = structured_list_count
        self.manual_list_count = manual_list_count

    @classmethod
    def from_path(cls, docx_path: Path) -> _DocxXmlContext:
        with ZipFile(docx_path) as archive:
            document_root = _read_xml(archive, "word/document.xml")
            styles_root = _read_xml(archive, "word/styles.xml")
            settings_root = _read_xml(archive, "word/settings.xml")
            core_root = _read_xml(archive, "docProps/core.xml")

        return cls(
            languages=_collect_languages(styles_root, settings_root, core_root),
            image_alt_values=_collect_image_alt_values(document_root),
            links=_collect_hyperlink_texts(document_root),
            table_header_count=_count_xml_table_headers(document_root),
            structured_list_count=_count_structured_lists(document_root),
            manual_list_count=0,
        )


class _DocxContext:
    def __init__(
        self,
        *,
        resource_title: str,
        filename_stem: str,
        title: str | None,
        text: str,
        languages: list[str],
        heading_levels: list[int],
        image_alt_values: list[str | None],
        table_count: int,
        table_header_count: int,
        links: list[str],
        structured_list_count: int,
        manual_list_count: int,
    ) -> None:
        self.resource_title = resource_title
        self.filename_stem = filename_stem
        self.title = title
        self.text = text
        self.languages = languages
        self.heading_levels = heading_levels
        self.image_alt_values = image_alt_values
        self.table_count = table_count
        self.table_header_count = table_header_count
        self.links = links
        self.structured_list_count = structured_list_count
        self.manual_list_count = manual_list_count

    @classmethod
    def from_document(
        cls,
        resource: Any,
        docx_path: Path,
        document: Any,
        xml_context: _DocxXmlContext,
    ) -> _DocxContext:
        paragraphs = list(document.paragraphs)
        text = _normalize_text(" ".join(paragraph.text for paragraph in paragraphs if paragraph.text))
        manual_list_count = sum(
            1
            for paragraph in paragraphs
            if MANUAL_LIST_RE.match(paragraph.text or "") and not _paragraph_has_numbering(paragraph)
        )
        structured_list_count = xml_context.structured_list_count or sum(
            1 for paragraph in paragraphs if _paragraph_has_numbering(paragraph)
        )
        return cls(
            resource_title=_resource_title(resource),
            filename_stem=docx_path.stem,
            title=_core_title(document),
            text=text,
            languages=xml_context.languages,
            heading_levels=_heading_levels(paragraphs),
            image_alt_values=xml_context.image_alt_values,
            table_count=len(document.tables),
            table_header_count=xml_context.table_header_count or _count_identifiable_table_headers(document),
            links=xml_context.links,
            structured_list_count=structured_list_count,
            manual_list_count=manual_list_count,
        )


def _check_readable() -> AccessibilityCheckResult:
    return _result(
        "docx.readable",
        "Documento legible",
        "PASS",
        "El DOCX se puede abrir y parsear.",
        "Mantén el documento en formato Word válido y no corrupto.",
    )


def _check_extractable_text(context: _DocxContext) -> AccessibilityCheckResult:
    text_length = len(context.text)
    if text_length >= 40:
        return _result(
            "docx.extractable_text",
            "Texto extraíble",
            "PASS",
            f"Se han extraído {text_length} caracteres de texto.",
            "Verifica igualmente que el orden de lectura sea lógico.",
            "WCAG 1.3.2",
        )
    if text_length > 0:
        return _result(
            "docx.extractable_text",
            "Texto extraíble",
            "WARNING",
            f"Solo se han extraído {text_length} caracteres de texto.",
            "Revisa si el documento contiene suficiente texto real y no solo imágenes.",
            "WCAG 1.3.2",
        )
    return _result(
        "docx.extractable_text",
        "Texto extraíble",
        "FAIL",
        "No se ha extraído texto útil del DOCX.",
        "Incluye texto editable o proporciona una versión accesible del contenido.",
        "WCAG 1.3.2",
    )


def _check_language(context: _DocxContext) -> AccessibilityCheckResult:
    if context.languages:
        return _result(
            "docx.lang",
            "Idioma del documento",
            "PASS",
            f"Se detecta idioma en el documento: {', '.join(context.languages[:3])}.",
            "Mantén definido el idioma principal del documento.",
            "WCAG 3.1.1",
        )
    return _result(
        "docx.lang",
        "Idioma del documento",
        "FAIL",
        "No se detecta idioma en estilos, docDefaults ni propiedades.",
        "Define el idioma principal del documento en Word.",
        "WCAG 3.1.1",
    )


def _check_title(context: _DocxContext) -> AccessibilityCheckResult:
    title = _normalize_text(context.title or "")
    if title and not _is_generic_docx_title(title, context):
        return _result(
            "docx.title",
            "Título del documento",
            "PASS",
            f"El DOCX tiene título de propiedades: \"{_shorten(title)}\".",
            "Usa títulos únicos y descriptivos en las propiedades del documento.",
            "WCAG 2.4.2",
        )
    return _result(
        "docx.title",
        "Título del documento",
        "WARNING",
        "No hay título útil en las propiedades del documento.",
        "Añade un título descriptivo en Archivo > Información > Propiedades.",
        "WCAG 2.4.2",
    )


def _check_heading_styles(context: _DocxContext) -> AccessibilityCheckResult:
    if context.heading_levels:
        levels = ", ".join(f"Heading {level}" for level in sorted(set(context.heading_levels)))
        return _result(
            "docx.headings",
            "Uso de estilos de encabezado",
            "PASS",
            f"Se detectan estilos semánticos de encabezado: {levels}.",
            "Usa estilos Heading para estructurar secciones, no solo formato visual.",
            "WCAG 1.3.1",
        )
    return _result(
        "docx.headings",
        "Uso de estilos de encabezado",
        "WARNING",
        "No se detectan estilos Heading 1, Heading 2, etc.",
        "Aplica estilos de encabezado de Word a los títulos de sección.",
        "WCAG 1.3.1",
    )


def _check_heading_hierarchy(context: _DocxContext) -> AccessibilityCheckResult:
    if not context.heading_levels:
        return _result(
            "docx.heading_hierarchy",
            "Jerarquía de encabezados",
            "NOT_APPLICABLE",
            "No hay encabezados semánticos para evaluar jerarquía.",
            "Primero aplica estilos Heading y después revisa la jerarquía.",
            "WCAG 1.3.1",
        )
    skips: list[str] = []
    previous = context.heading_levels[0]
    for level in context.heading_levels[1:]:
        if level - previous > 1:
            skips.append(f"Heading {previous} -> Heading {level}")
        previous = level
    if not skips:
        return _result(
            "docx.heading_hierarchy",
            "Jerarquía de encabezados",
            "PASS",
            "La secuencia de encabezados no presenta saltos bruscos.",
            "Conserva una jerarquía de encabezados ordenada.",
            "WCAG 1.3.1",
        )
    return _result(
        "docx.heading_hierarchy",
        "Jerarquía de encabezados",
        "WARNING",
        f"Se detectan saltos de jerarquía: {'; '.join(skips[:3])}.",
        "Inserta los niveles intermedios necesarios o ajusta los estilos de encabezado.",
        "WCAG 1.3.1",
    )


def _check_image_alt(context: _DocxContext) -> AccessibilityCheckResult:
    if not context.image_alt_values:
        return _result(
            "docx.image_alt",
            "Imágenes con texto alternativo",
            "NOT_APPLICABLE",
            "No se detectan imágenes en el documento.",
            "Cuando añadas imágenes informativas, incluye texto alternativo.",
            "WCAG 1.1.1",
        )
    missing = [value for value in context.image_alt_values if not value]
    if missing:
        return _result(
            "docx.image_alt",
            "Imágenes con texto alternativo",
            "FAIL",
            f"{len(missing)} imagen(es) no tienen title/descr/alt en el XML.",
            "Añade texto alternativo a las imágenes informativas desde Word.",
            "WCAG 1.1.1",
        )
    return _result(
        "docx.image_alt",
        "Imágenes con texto alternativo",
        "PASS",
        "Todas las imágenes detectadas tienen texto alternativo.",
        "Mantén alternativas equivalentes al propósito de cada imagen.",
        "WCAG 1.1.1",
    )


def _check_tables(context: _DocxContext) -> AccessibilityCheckResult:
    if context.table_count == 0:
        return _result(
            "docx.tables",
            "Tablas con estructura",
            "NOT_APPLICABLE",
            "No se detectan tablas.",
            "Si añades tablas de datos, marca o identifica la fila de encabezado.",
            "WCAG 1.3.1",
        )
    if context.table_header_count >= context.table_count:
        return _result(
            "docx.tables",
            "Tablas con estructura",
            "PASS",
            f"{context.table_count} tabla(s) tienen primera fila marcada o identificable como encabezado.",
            "Verifica manualmente relaciones complejas entre encabezados y celdas.",
            "WCAG 1.3.1",
        )
    return _result(
        "docx.tables",
        "Tablas con estructura",
        "WARNING",
        f"{context.table_count - context.table_header_count} tabla(s) no tienen encabezado identificable.",
        "Marca la primera fila como encabezado o usa una estructura de tabla clara.",
        "WCAG 1.3.1",
    )


def _check_links(context: _DocxContext) -> AccessibilityCheckResult:
    if not context.links:
        return _result(
            "docx.links",
            "Enlaces descriptivos",
            "NOT_APPLICABLE",
            "No se detectan enlaces.",
            "Cuando añadas enlaces, usa textos que indiquen destino o acción.",
            "WCAG 2.4.4",
        )
    problematic = [
        link for link in context.links if _is_generic_link_text(link) or URL_TEXT_RE.match(link)
    ]
    if problematic:
        return _result(
            "docx.links",
            "Enlaces descriptivos",
            "FAIL",
            f"Hay enlaces con texto poco descriptivo: {', '.join(_shorten(item) for item in problematic[:4])}.",
            "Sustituye textos como “aquí” o URLs desnudas por descripciones del destino.",
            "WCAG 2.4.4",
        )
    return _result(
        "docx.links",
        "Enlaces descriptivos",
        "PASS",
        "Los enlaces detectados tienen textos descriptivos.",
        "Mantén textos de enlace comprensibles fuera de contexto.",
        "WCAG 2.4.4",
    )


def _check_lists(context: _DocxContext) -> AccessibilityCheckResult:
    if context.structured_list_count > 0:
        return _result(
            "docx.lists",
            "Listas estructuradas",
            "PASS",
            f"Se detectan {context.structured_list_count} párrafo(s) con numeración/viñetas de Word.",
            "Usa siempre listas nativas de Word para secuencias o grupos de elementos.",
            "WCAG 1.3.1",
        )
    if context.manual_list_count > 0:
        return _result(
            "docx.lists",
            "Listas estructuradas",
            "WARNING",
            f"{context.manual_list_count} línea(s) parecen listas manuales sin estructura de Word.",
            "Convierte guiones o numeración escrita a listas nativas de Word.",
            "WCAG 1.3.1",
        )
    return _result(
        "docx.lists",
        "Listas estructuradas",
        "NOT_APPLICABLE",
        "No se detectan listas.",
        "Cuando añadas listas, usa viñetas o numeración de Word.",
        "WCAG 1.3.1",
    )


def _should_attempt_docx_scan(resource: Any, core: Any) -> bool:
    if core.origin in {"RALTI", "LTI", "EXTERNAL_URL"}:
        return False
    if core.accessStatus in {"REQUIERE_SSO", "REQUIERE_INTERACCION", "NO_ANALIZABLE"}:
        return False
    if not core.contentAvailable:
        return False
    return core.type == "DOCX" or _resource_has_docx_reference(resource)


def _content_is_docx(content: ResourceContentResult) -> bool:
    mime_type = (content.mimeType or "").split(";", 1)[0].strip().lower()
    filename = (content.filename or content.binaryPath or content.title or "").lower()
    return mime_type == DOCX_MIME_TYPE or filename.endswith(".docx")


def _resource_has_docx_reference(resource: Any) -> bool:
    if isinstance(resource, dict):
        values = [
            resource.get("localPath"),
            resource.get("filePath"),
            resource.get("path"),
            resource.get("sourceUrl"),
            resource.get("downloadUrl"),
            resource.get("url"),
            resource.get("title"),
            resource.get("filename"),
            resource.get("mimeType"),
            resource.get("contentType"),
        ]
    else:
        values = [
            getattr(resource, "localPath", None),
            getattr(resource, "file_path", None),
            getattr(resource, "path", None),
            getattr(resource, "source_url", None),
            getattr(resource, "download_url", None),
            getattr(resource, "url", None),
            getattr(resource, "title", None),
            getattr(resource, "filename", None),
            getattr(resource, "mimeType", None),
            getattr(resource, "content_type", None),
        ]
    return any(str(value).lower().endswith(".docx") or str(value).lower() == DOCX_MIME_TYPE for value in values if value)


def _read_xml(archive: ZipFile, name: str) -> ElementTree.Element | None:
    try:
        return ElementTree.fromstring(archive.read(name))
    except (KeyError, ElementTree.ParseError):
        return None


def _collect_languages(*roots: ElementTree.Element | None) -> list[str]:
    languages: list[str] = []
    for root in roots:
        if root is None:
            continue
        for element in root.iter():
            tag = _local_name(element.tag)
            if tag == "lang":
                for key, value in element.attrib.items():
                    if _local_name(key) in {"val", "bidi", "eastAsia"} and value:
                        languages.append(value)
            elif tag == "language" and element.text:
                languages.append(element.text.strip())
    return sorted(set(filter(None, languages)))


def _collect_image_alt_values(root: ElementTree.Element | None) -> list[str | None]:
    if root is None:
        return []
    values: list[str | None] = []
    for element in root.iter():
        if _local_name(element.tag) != "docPr":
            continue
        alt_text = _normalize_text(element.attrib.get("descr") or element.attrib.get("title") or "")
        values.append(alt_text or None)
    return values


def _collect_hyperlink_texts(root: ElementTree.Element | None) -> list[str]:
    if root is None:
        return []
    links: list[str] = []
    for element in root.iter():
        if _local_name(element.tag) != "hyperlink":
            continue
        text = _normalize_text(" ".join(child.text or "" for child in element.iter() if _local_name(child.tag) == "t"))
        if text:
            links.append(text)
    return links


def _count_xml_table_headers(root: ElementTree.Element | None) -> int:
    if root is None:
        return 0
    count = 0
    for table in (element for element in root.iter() if _local_name(element.tag) == "tbl"):
        if any(_local_name(element.tag) == "tblHeader" for element in table.iter()):
            count += 1
            continue
        rows = [element for element in table if _local_name(element.tag) == "tr"]
        if rows and _row_has_text(rows[0]):
            count += 1
    return count


def _count_structured_lists(root: ElementTree.Element | None) -> int:
    if root is None:
        return 0
    return sum(1 for element in root.iter() if _local_name(element.tag) == "numPr")


def _count_identifiable_table_headers(document: Any) -> int:
    count = 0
    for table in document.tables:
        if not table.rows:
            continue
        first_row_text = [_normalize_text(cell.text) for cell in table.rows[0].cells]
        if any(first_row_text):
            count += 1
    return count


def _row_has_text(row: ElementTree.Element) -> bool:
    return any((element.text or "").strip() for element in row.iter() if _local_name(element.tag) == "t")


def _paragraph_has_numbering(paragraph: Any) -> bool:
    p_pr = getattr(paragraph._p, "pPr", None)
    return bool(p_pr is not None and getattr(p_pr, "numPr", None) is not None)


def _heading_levels(paragraphs: list[Any]) -> list[int]:
    levels: list[int] = []
    for paragraph in paragraphs:
        style = paragraph.style
        style_name = (getattr(style, "name", "") or "").strip().lower()
        style_id = (getattr(style, "style_id", "") or "").strip().lower()
        match = re.search(r"heading\s*(\d+)", style_name) or re.search(r"heading(\d+)", style_id)
        if match:
            levels.append(int(match.group(1)))
    return levels


def _core_title(document: Any) -> str | None:
    title = getattr(document.core_properties, "title", None)
    return str(title) if title else None


def _is_generic_docx_title(title: str, context: _DocxContext) -> bool:
    normalized = _normalize_text(title).strip(" .:-").lower()
    filename = _normalize_text(context.filename_stem).strip(" .:-").lower()
    resource_title = _normalize_text(context.resource_title).strip(" .:-").lower()
    return normalized in GENERIC_DOCX_TITLES or normalized in {filename, resource_title}


def _is_generic_link_text(value: str) -> bool:
    normalized = _normalize_text(value).strip(" .:-").lower()
    return normalized in GENERIC_LINK_TEXTS


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _resource_title(resource: Any) -> str:
    if isinstance(resource, dict):
        return _normalize_text(str(resource.get("title") or ""))
    return _normalize_text(str(getattr(resource, "title", "") or ""))


def _result(
    check_id: str,
    title: str,
    status: str,
    evidence: str,
    recommendation: str,
    wcag_hint: str | None = None,
) -> AccessibilityCheckResult:
    return AccessibilityCheckResult(
        checkId=check_id,
        checkTitle=title,
        status=status,
        evidence=evidence,
        recommendation=recommendation,
        wcagHint=wcag_hint,
    )


def _normalize_text(value: str) -> str:
    return " ".join(value.strip().split())


def _shorten(value: str, limit: int = 80) -> str:
    normalized = _normalize_text(value)
    return normalized if len(normalized) <= limit else f"{normalized[: limit - 1]}…"
