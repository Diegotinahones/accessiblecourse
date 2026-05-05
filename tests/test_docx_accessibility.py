from __future__ import annotations

from pathlib import Path

from docx import Document

from app.services.docx_accessibility import analyze_docx_accessibility, run_docx_accessibility_scan
from app.services.storage import get_extracted_dir

TINY_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff?"
    b"\x00\x05\xfe\x02\xfeA\x0d\x89\xb1\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _write_docx(path: Path, *, title: str | None = "Guia accesible", headings: tuple[int, ...] = ()) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    if title is not None:
        document.core_properties.title = title
    document.add_paragraph(
        "Este documento Word contiene texto suficiente para validar la extraccion automatica "
        "de contenido y preparar los checks de accesibilidad."
    )
    for level in headings:
        document.add_heading(f"Encabezado nivel {level}", level=level)
    document.save(path)


def _status_by_check(path: Path) -> dict[str, str]:
    checks = analyze_docx_accessibility({"id": "docx", "title": "Guia accesible", "type": "DOCX"}, path)
    return {check.checkId: check.status for check in checks}


def test_docx_accessibility_passes_extractable_text(tmp_path) -> None:
    docx_path = tmp_path / "guide.docx"
    _write_docx(docx_path)

    statuses = _status_by_check(docx_path)

    assert statuses["docx.readable"] == "PASS"
    assert statuses["docx.extractable_text"] == "PASS"


def test_docx_accessibility_warns_missing_title(tmp_path) -> None:
    docx_path = tmp_path / "guide.docx"
    _write_docx(docx_path, title=None)

    statuses = _status_by_check(docx_path)

    assert statuses["docx.title"] == "WARNING"


def test_docx_accessibility_passes_heading_styles(tmp_path) -> None:
    docx_path = tmp_path / "headings.docx"
    _write_docx(docx_path, headings=(1, 2))

    statuses = _status_by_check(docx_path)

    assert statuses["docx.headings"] == "PASS"
    assert statuses["docx.heading_hierarchy"] == "PASS"


def test_docx_accessibility_warns_heading_hierarchy_skip(tmp_path) -> None:
    docx_path = tmp_path / "heading-skip.docx"
    _write_docx(docx_path, headings=(1, 3))

    statuses = _status_by_check(docx_path)

    assert statuses["docx.heading_hierarchy"] == "WARNING"


def test_docx_accessibility_fails_image_without_alt(tmp_path) -> None:
    image_path = tmp_path / "pixel.png"
    image_path.write_bytes(TINY_PNG)
    docx_path = tmp_path / "image.docx"
    document = Document()
    document.core_properties.title = "Documento con imagen"
    document.add_paragraph("Documento con una imagen informativa sin texto alternativo suficiente.")
    document.add_picture(str(image_path))
    document.save(docx_path)

    statuses = _status_by_check(docx_path)

    assert statuses["docx.image_alt"] == "FAIL"


def test_docx_accessibility_job_scan_treats_other_docx_as_docx(test_settings) -> None:
    job_id = "44444444-4444-4444-4444-444444444444"
    docx_path = get_extracted_dir(test_settings, job_id) / "docs" / "guide.docx"
    _write_docx(docx_path)

    report = run_docx_accessibility_scan(
        settings=test_settings,
        job_id=job_id,
        resources=[
            {
                "id": "word-guide",
                "title": "Guia Word",
                "type": "OTHER",
                "origin": "INTERNAL_FILE",
                "localPath": "docs/guide.docx",
                "accessStatus": "OK",
                "canAccess": True,
                "canDownload": True,
                "contentAvailable": True,
            }
        ],
    )

    assert report.summary.docxResourcesTotal == 1
    assert report.summary.docxResourcesAnalyzed == 1
    assert report.modules[0].resources[0].analysisType == "DOCX"


def test_docx_accessibility_job_scan_records_unexpected_resource_error(test_settings, monkeypatch) -> None:
    job_id = "66666666-6666-6666-6666-666666666666"
    docx_path = get_extracted_dir(test_settings, job_id) / "docs" / "broken-analysis.docx"
    _write_docx(docx_path)

    def raise_unexpected_error(*args, **kwargs):
        raise RuntimeError("boom")

    monkeypatch.setattr(
        "app.services.docx_accessibility.analyze_docx_accessibility",
        raise_unexpected_error,
    )

    report = run_docx_accessibility_scan(
        settings=test_settings,
        job_id=job_id,
        resources=[
            {
                "id": "word-guide",
                "title": "Guia Word",
                "type": "DOCX",
                "origin": "INTERNAL_FILE",
                "localPath": "docs/broken-analysis.docx",
                "accessStatus": "OK",
                "canAccess": True,
                "canDownload": True,
                "contentAvailable": True,
            }
        ],
    )

    checks = report.modules[0].resources[0].checks
    assert report.summary.docxResourcesTotal == 1
    assert checks[0].checkId == "docx.analysis"
    assert checks[0].status == "ERROR"
    assert "RuntimeError" in checks[0].evidence


def test_docx_accessibility_job_scan_skips_non_docx_resources(test_settings) -> None:
    report = run_docx_accessibility_scan(
        settings=test_settings,
        job_id="55555555-5555-5555-5555-555555555555",
        resources=[
            {
                "id": "pdf-guide",
                "title": "Guia PDF",
                "type": "PDF",
                "origin": "INTERNAL_FILE",
                "localPath": "docs/guide.pdf",
                "accessStatus": "OK",
                "contentAvailable": True,
            }
        ],
    )

    assert report.summary.docxResourcesTotal == 0
    assert report.modules == []
