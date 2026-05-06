from __future__ import annotations

import io
import json
from datetime import datetime, timezone
from docx import Document
from pathlib import Path
from reportlab.pdfgen import canvas as pdf_canvas
from zipfile import ZipFile

from sqlalchemy.exc import IntegrityError
from sqlmodel import Session, select

from app.models import Job as ProcessingJob, JobEvent
from app.models.entities import ChecklistResponse
from app.services.resource_core import get_resource_content
from app.services.url_check import UrlCheckResult
from tests.conftest import build_sample_imscc

DOCX_TINY_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff?"
    b"\x00\x05\xfe\x02\xfeA\x0d\x89\xb1\x00\x00\x00\x00IEND\xaeB`\x82"
)


def write_inventory(test_settings, job_id: str, resources: list[dict[str, object]]) -> None:
    job_dir = test_settings.storage_root / "jobs" / job_id
    job_dir.mkdir(parents=True, exist_ok=True)
    (job_dir / "resources.json").write_text(json.dumps(resources), encoding="utf-8")


def build_large_imscc(payload_size: int) -> bytes:
    buffer = io.BytesIO()
    with ZipFile(buffer, "w") as archive:
        archive.writestr(
            "imsmanifest.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<manifest xmlns="http://www.imsglobal.org/xsd/imscp_v1p1">
  <resources>
    <resource identifier="res-1" type="webcontent" href="module_1/large.bin">
      <file href="module_1/large.bin" />
    </resource>
  </resources>
</manifest>
""",
        )
        archive.writestr("module_1/large.bin", b"A" * payload_size)
    return buffer.getvalue()


def build_text_pdf(title: str = "Brief de la PEC") -> bytes:
    buffer = io.BytesIO()
    document = pdf_canvas.Canvas(buffer)
    document.drawString(72, 760, title)
    document.drawString(72, 735, "Documento de prueba con texto seleccionable.")
    document.save()
    return buffer.getvalue()


def build_text_docx(title: str = "Ficha Word") -> bytes:
    buffer = io.BytesIO()
    document = Document()
    document.core_properties.title = title
    document.add_paragraph(
        "Este documento Word contiene texto suficiente para validar el analisis automatico "
        "sin incluir binarios ni contenido completo en el informe."
    )
    document.add_picture(io.BytesIO(DOCX_TINY_PNG))
    document.save(buffer)
    return buffer.getvalue()


def test_offline_processing_marks_error_if_persistence_fails(client, monkeypatch) -> None:
    def fail_after_save_progress(session, settings, **kwargs) -> None:
        try:
            session.connection().exec_driver_sql("INSERT INTO job (id) VALUES (?)", ("broken-job",))
        except IntegrityError as exc:
            raise RuntimeError("forced persistence failure") from exc

    monkeypatch.setattr("app.services.jobs._persist_analyzed_inventory", fail_after_save_progress)

    create_response = client.post(
        "/api/jobs",
        files={"file": ("course.imscc", build_sample_imscc(), "application/octet-stream")},
    )

    assert create_response.status_code == 201, create_response.text
    job_id = create_response.json()["jobId"]

    status_response = client.get(f"/api/jobs/{job_id}")
    assert status_response.status_code == 200, status_response.text
    payload = status_response.json()
    assert payload["status"] == "error"
    assert payload["phase"] == "ERROR"
    assert payload["progress"] == 95
    assert payload["errorCode"] == "unexpected_processing_error"


def build_imscc_with_external_link() -> bytes:
    buffer = io.BytesIO()
    with ZipFile(buffer, "w") as archive:
        archive.writestr(
            "imsmanifest.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<manifest xmlns="http://www.imsglobal.org/xsd/imsccv1p1/imscp_v1p1">
  <organizations>
    <organization identifier="org-1">
      <title>Curso Demo</title>
      <item identifier="module-1">
        <title>Tema 1</title>
        <item identifier="item-1" identifierref="res-pdf">
          <title>Guía</title>
        </item>
        <item identifier="item-2" identifierref="res-link">
          <title>Enlace externo</title>
        </item>
        <item identifier="item-3" identifierref="res-metadata">
          <title>Metadata</title>
        </item>
      </item>
    </organization>
  </organizations>
  <resources>
    <resource identifier="res-pdf" type="webcontent" href="course/topic-1/guide.pdf">
      <file href="course/topic-1/guide.pdf" />
    </resource>
    <resource identifier="res-link" type="imswl_xmlv1p1" href="web_resources/external_link.xml">
      <file href="web_resources/external_link.xml" />
    </resource>
    <resource identifier="res-metadata" type="webcontent" href="metadata/descriptor.xml">
      <file href="metadata/descriptor.xml" />
    </resource>
  </resources>
</manifest>
""",
        )
        archive.writestr("course/topic-1/guide.pdf", b"%PDF-1.4\n%offline pdf\n")
        archive.writestr(
            "web_resources/external_link.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<webLink xmlns="http://www.imsglobal.org/xsd/imswl_v1p1">
  <title>Enlace externo</title>
  <url href="https://example.com/broken-link" />
</webLink>
""",
        )
        archive.writestr("metadata/descriptor.xml", "<metadata />")
    return buffer.getvalue()


def build_imscc_with_lti_noise() -> bytes:
    buffer = io.BytesIO()
    with ZipFile(buffer, "w") as archive:
        archive.writestr(
            "imsmanifest.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<manifest xmlns="http://www.imsglobal.org/xsd/imsccv1p1/imscp_v1p1">
  <organizations>
    <organization identifier="org-1">
      <title>Curso limpio</title>
      <item identifier="module-1">
        <title>Unidad principal</title>
        <item identifier="item-1" identifierref="res-pdf">
          <title>Guía principal</title>
        </item>
        <item identifier="item-2" identifierref="res-lti-1">
          <title>Learning Tools</title>
        </item>
        <item identifier="item-3" identifierref="res-lti-2">
          <title>Recursos de Aprendizaje</title>
        </item>
        <item identifier="item-4" identifierref="res-meta">
          <title>Metadata</title>
        </item>
      </item>
    </organization>
  </organizations>
  <resources>
    <resource identifier="res-pdf" type="webcontent" href="course/guide.pdf">
      <file href="course/guide.pdf" />
    </resource>
    <resource identifier="res-lti-1" type="imsbasiclti_xmlv1p0" href="lti_resource_links/tool-1.xml">
      <file href="lti_resource_links/tool-1.xml" />
    </resource>
    <resource identifier="res-lti-2" type="imsbasiclti_xmlv1p0" href="lti_resource_links/tool-2.xml">
      <file href="lti_resource_links/tool-2.xml" />
    </resource>
    <resource identifier="res-meta" type="webcontent" href="metadata/descriptor.xml">
      <file href="metadata/descriptor.xml" />
    </resource>
  </resources>
</manifest>
""",
        )
        archive.writestr("course/guide.pdf", b"%PDF-1.4\n%clean guide\n")
        archive.writestr(
            "lti_resource_links/tool-1.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<webLink xmlns="http://www.imsglobal.org/xsd/imswl_v1p1">
  <title>Learning Tools</title>
  <url href="https://ralti.uoc.edu/launch?tool=abc" />
</webLink>
""",
        )
        archive.writestr(
            "lti_resource_links/tool-2.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<webLink xmlns="http://www.imsglobal.org/xsd/imswl_v1p1">
  <title>Recursos de Aprendizaje</title>
  <url href="https://ralti.uoc.edu/launch?tool=abc" />
</webLink>
""",
        )
        archive.writestr("metadata/descriptor.xml", "<metadata />")
    return buffer.getvalue()


def build_imscc_with_unmapped_resource() -> bytes:
    buffer = io.BytesIO()
    with ZipFile(buffer, "w") as archive:
        archive.writestr(
            "imsmanifest.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<manifest xmlns="http://www.imsglobal.org/xsd/imsccv1p1/imscp_v1p1">
  <organizations>
    <organization identifier="org-1">
      <title>Curso Demo</title>
      <item identifier="module-1">
        <title>Bloque 1</title>
        <item identifier="item-1" identifierref="res-guide">
          <title>Guía principal</title>
        </item>
      </item>
    </organization>
  </organizations>
  <resources>
    <resource identifier="res-guide" type="webcontent" href="course/guide.pdf">
      <file href="course/guide.pdf" />
    </resource>
    <resource identifier="res-orphan" type="webcontent" href="web_resources/orphan.html">
      <file href="web_resources/orphan.html" />
    </resource>
  </resources>
</manifest>
""",
        )
        archive.writestr("course/guide.pdf", b"%PDF-1.4\n%offline pdf\n")
        archive.writestr("web_resources/orphan.html", "<html><body>Orphan</body></html>")
    return buffer.getvalue()


def build_imscc_with_offline_deep_scan() -> bytes:
    buffer = io.BytesIO()
    with ZipFile(buffer, "w") as archive:
        archive.writestr(
            "imsmanifest.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<manifest xmlns="http://www.imsglobal.org/xsd/imsccv1p1/imscp_v1p1">
  <organizations>
    <organization identifier="org-1">
      <title>Curso con deep scan</title>
      <item identifier="module-1">
        <title>Unidad 1</title>
        <item identifier="item-1" identifierref="res-html">
          <title>Página principal</title>
        </item>
      </item>
    </organization>
  </organizations>
  <resources>
    <resource identifier="res-html" type="webcontent" href="course/module/page.html">
      <file href="course/module/page.html" />
    </resource>
  </resources>
</manifest>
""",
        )
        archive.writestr(
            "course/module/page.html",
            """
            <html><body>
              <a href="../downloads/worksheet.docx">Ficha descargable</a>
              <a href="../downloads/worksheet.docx#dup">Ficha descargable</a>
              <img src="../images/chart.png" alt="Diagrama del curso" />
              <a href="nested/page2.html">Más materiales</a>
              <a href="https://example.com/resource">Recurso web</a>
              <iframe src="https://www.youtube.com/watch?v=demo123"></iframe>
            </body></html>
            """,
        )
        archive.writestr(
            "course/module/nested/page2.html",
            """
            <html><body>
              <a href="../../downloads/slides.pptx">Presentación final</a>
            </body></html>
            """,
        )
        archive.writestr("course/downloads/worksheet.docx", b"DOCX")
        archive.writestr("course/downloads/slides.pptx", b"PPTX")
        archive.writestr("course/images/chart.png", b"PNG")
    return buffer.getvalue()


def build_imscc_with_pec_section_deep_scan(*, include_video: bool = False) -> bytes:
    buffer = io.BytesIO()
    with ZipFile(buffer, "w") as archive:
        archive.writestr(
            "imsmanifest.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<manifest xmlns="http://www.imsglobal.org/xsd/imsccv1p1/imscp_v1p1">
  <organizations>
    <organization identifier="org-1">
      <title>Curso PEC</title>
      <item identifier="section-pec1">
        <title>PEC 1: Actividad inicial</title>
        <item identifier="item-html" identifierref="res-html">
          <title>Página de actividad</title>
        </item>
      </item>
    </organization>
  </organizations>
  <resources>
    <resource identifier="res-html" type="webcontent" href="course/pec1/activity.html">
      <file href="course/pec1/activity.html" />
    </resource>
  </resources>
</manifest>
""",
        )
        video_link = (
            '<a href="https://www.youtube.com/watch?v=demo-pec">Vídeo de apoyo</a>'
            if include_video
            else ""
        )
        archive.writestr(
            "course/pec1/activity.html",
            f"""
            <html><body>
              <a href="../downloads/brief.pdf">Brief de la PEC</a>
              <a href="../downloads/ficha.docx">Ficha Word</a>
              {video_link}
              <a href="https://example.com/como-citar-ia">¿Cómo citar la IA…</a>
            </body></html>
            """,
        )
        archive.writestr("course/downloads/brief.pdf", build_text_pdf())
        archive.writestr("course/downloads/ficha.docx", build_text_docx())
    return buffer.getvalue()


def collect_resource_ids(node: dict[str, object]) -> list[str]:
    collected: list[str] = []
    resource_id = node.get("resourceId")
    if isinstance(resource_id, str) and resource_id:
        collected.append(resource_id)
    children = node.get("children")
    if isinstance(children, list):
        for child in children:
            if isinstance(child, dict):
                collected.extend(collect_resource_ids(child))
    return collected


def test_bootstrap_inventory_and_persist_checklist(client, test_settings) -> None:
    job_id = "job-thread4-bootstrap"
    write_inventory(
        test_settings,
        job_id,
        [
            {
                "id": "web-home",
                "title": "Portada del curso",
                "type": "WEB",
                "origin": "imscc:webcontent",
                "url": "https://example.edu/course/home",
                "path": "course/home.html",
                "course_path": "Inicio/Portada",
                "status": "OK",
            },
            {
                "id": "pdf-guide",
                "title": "Guia docente",
                "type": "PDF",
                "origin": "imscc:file",
                "url": "https://example.edu/files/guide.pdf",
                "path": "docs/guide.pdf",
                "course_path": "Documentacion/Guia",
                "status": "WARN",
                "notes": ["PDF pendiente de etiquetado"],
            },
        ],
    )

    resources_response = client.get(f"/api/jobs/{job_id}/resources")
    assert resources_response.status_code == 200, resources_response.text
    resources_payload = resources_response.json()
    assert resources_payload["jobId"] == job_id
    assert len(resources_payload["resources"]) == 2
    assert resources_payload["reviewSession"]["status"] == "NOT_STARTED"

    detail_response = client.get(f"/api/jobs/{job_id}/resources/pdf-guide")
    assert detail_response.status_code == 200, detail_response.text
    detail_payload = detail_response.json()
    assert detail_payload["resource"]["notes"] == "PDF pendiente de etiquetado"
    assert all(item["value"] == "PENDING" for item in detail_payload["checklist"]["items"])

    save_response = client.put(
        f"/api/jobs/{job_id}/resources/pdf-guide/checklist",
        json={
            "responses": [
                {
                    "itemKey": "tagged",
                    "value": "FAIL",
                    "comment": "No tiene estructura etiquetada.",
                },
                {"itemKey": "lang", "value": "PASS"},
            ]
        },
    )
    assert save_response.status_code == 200, save_response.text
    save_payload = save_response.json()
    assert save_payload["resourceId"] == "pdf-guide"
    assert save_payload["reviewState"] == "NEEDS_FIX"
    assert save_payload["failCount"] == 1

    persisted_detail_response = client.get(f"/api/jobs/{job_id}/resources/pdf-guide")
    assert persisted_detail_response.status_code == 200, persisted_detail_response.text
    persisted_detail = persisted_detail_response.json()
    tagged_item = next(item for item in persisted_detail["checklist"]["items"] if item["itemKey"] == "tagged")
    assert tagged_item["value"] == "FAIL"
    assert tagged_item["comment"] == "No tiene estructura etiquetada."
    assert persisted_detail["resource"]["reviewState"] == "NEEDS_FIX"
    assert persisted_detail["reviewSession"]["status"] == "IN_PROGRESS"

    summary_response = client.get(f"/api/jobs/{job_id}/summary")
    assert summary_response.status_code == 200, summary_response.text
    summary_payload = summary_response.json()
    assert summary_payload["totalResources"] == 2
    assert summary_payload["totalFailItems"] == 1
    assert summary_payload["resources"][0]["resourceId"] == "pdf-guide"
    assert summary_payload["resources"][0]["recommendations"][0]["itemKey"] == "tagged"
    assert summary_payload["resources"][0]["recommendations"][0]["recommendation"]

    report_response = client.post(f"/api/jobs/{job_id}/report")
    assert report_response.status_code == 200, report_response.text
    report_payload = report_response.json()
    assert report_payload["meta"]["jobId"] == job_id
    assert report_payload["stats"] == {"resources": 2, "fails": 1, "pending": 14}
    assert report_payload["summary"]["topResources"][0]["resourceId"] == "pdf-guide"
    assert report_payload["routes"][0]["resources"]
    assert report_payload["resources"][0]["pending"] or report_payload["resources"][0]["fails"]

    report_dir = test_settings.storage_root / "jobs" / job_id / "report"
    json_path = report_dir / "report.json"
    docx_path = report_dir / "report.docx"
    pdf_path = report_dir / "report.pdf"
    assert json_path.exists()
    assert docx_path.exists()
    assert pdf_path.exists()
    assert docx_path.stat().st_size > 0
    assert pdf_path.stat().st_size > 0

    stored_payload = json.loads(json_path.read_text(encoding="utf-8"))
    assert stored_payload["reportId"] == report_payload["reportId"]
    assert stored_payload["summary"]["fails"] == 1
    assert stored_payload["summary"]["pending"] == 14

    latest_report_response = client.get(f"/api/jobs/{job_id}/report")
    assert latest_report_response.status_code == 200, latest_report_response.text
    latest_report_payload = latest_report_response.json()
    assert latest_report_payload["reportId"] == report_payload["reportId"]

    download_docx_response = client.get(f"/api/jobs/{job_id}/report/download?format=docx")
    assert download_docx_response.status_code == 200, download_docx_response.text
    assert (
        download_docx_response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment;" in download_docx_response.headers["content-disposition"]
    assert f"AccessibleCourse_Report_{job_id}_" in download_docx_response.headers["content-disposition"]
    assert len(download_docx_response.content) > 0

    download_pdf_response = client.get(f"/api/jobs/{job_id}/report/download?format=pdf")
    assert download_pdf_response.status_code == 200, download_pdf_response.text
    assert download_pdf_response.headers["content-type"] == "application/pdf"
    assert len(download_pdf_response.content) > 0

    download_json_response = client.get(f"/api/jobs/{job_id}/report/download?format=json")
    assert download_json_response.status_code == 200, download_json_response.text
    assert download_json_response.json()["reportId"] == report_payload["reportId"]


def test_upload_returns_clear_413_when_limit_is_exceeded(client, test_settings) -> None:
    oversized_imscc = build_large_imscc(payload_size=(test_settings.max_upload_bytes + 1024))

    response = client.post(
        "/api/jobs",
        files={"file": ("oversized-course.imscc", oversized_imscc, "application/octet-stream")},
    )

    assert response.status_code == 413, response.text
    payload = response.json()
    assert payload["code"] == "UPLOAD_TOO_LARGE"
    assert payload["message"] == "El archivo supera el l\u00edmite configurado"
    assert payload["maxMB"] == test_settings.max_upload_mb
    assert payload["actualMB"] > test_settings.max_upload_mb

    with Session(client.app.state.engine) as session:
        jobs = session.exec(select(ProcessingJob)).all()

    assert jobs == []


def test_upload_within_limit_returns_201(client) -> None:
    response = client.post(
        "/api/jobs",
        files={"file": ("course.imscc", build_sample_imscc(), "application/octet-stream")},
    )

    assert response.status_code == 201, response.text
    payload = response.json()
    assert payload["jobId"]

    with Session(client.app.state.engine) as session:
        jobs = session.exec(select(ProcessingJob)).all()

    assert len(jobs) == 1


def test_offline_inventory_groups_by_module_and_filters_broken_links(client, monkeypatch) -> None:
    class StubURLChecker:
        def check(self, resources):
            checked_at = datetime(2026, 4, 14, 10, 30, tzinfo=timezone.utc)
            results: dict[str, UrlCheckResult] = {}
            for resource in resources:
                url = resource.get("sourceUrl") or resource.get("url")
                if not url:
                    continue
                results[str(resource["id"])] = UrlCheckResult(
                    url=str(url),
                    checked=True,
                    broken_link=True,
                    reason="404_not_found",
                    status_code=404,
                    url_status="404",
                    final_url=str(url),
                    checked_at=checked_at,
                )
            return results

    monkeypatch.setattr("app.services.jobs.build_url_checker", lambda settings: StubURLChecker())

    create_response = client.post(
        "/api/jobs",
        files={"file": ("course.imscc", build_imscc_with_external_link(), "application/octet-stream")},
    )

    assert create_response.status_code == 201, create_response.text
    job_id = create_response.json()["jobId"]

    status_response = client.get(f"/api/jobs/{job_id}")
    assert status_response.status_code == 200, status_response.text
    assert status_response.json()["phase"] == "DONE"

    resources_response = client.get(f"/api/jobs/{job_id}/resources")
    assert resources_response.status_code == 200, resources_response.text
    resources_payload = resources_response.json()
    resources = resources_payload["resources"]
    assert len(resources) == 2
    assert [resource["title"] for resource in resources] == ["Guía", "Enlace externo"]

    pdf_resource = next(resource for resource in resources if resource["title"] == "Guía")
    link_resource = next(resource for resource in resources if resource["title"] == "Enlace externo")

    assert resources_payload["structure"]["title"] == "Curso Demo"
    organization = resources_payload["structure"]["organizations"][0]
    assert organization["title"] == "Curso Demo"
    assert organization["children"][0]["title"] == "Tema 1"
    assert organization["children"][0]["children"][0]["resourceId"] == "res-pdf"
    assert organization["children"][0]["children"][1]["resourceId"] == "res-link"
    assert resources_payload["structure"]["unplacedResourceIds"] == []

    assert pdf_resource["modulePath"] == "Tema 1"
    assert pdf_resource["coursePath"] == "Tema 1"
    assert pdf_resource["itemPath"] == "Tema 1 > Guía"
    assert pdf_resource["filePath"] == "course/topic-1/guide.pdf"
    assert pdf_resource["sourceUrl"] is None
    assert pdf_resource["canAccess"] is True
    assert pdf_resource["accessStatus"] == "OK"
    assert pdf_resource["httpStatus"] == 200
    assert pdf_resource["accessStatusCode"] == 200
    assert pdf_resource["canDownload"] is True
    assert pdf_resource["downloadStatusCode"] == 200
    assert pdf_resource["errorMessage"] is None
    assert pdf_resource["origin"] == "INTERNAL_FILE"
    assert pdf_resource["contentAvailable"] is True
    assert pdf_resource["localPath"] == "course/topic-1/guide.pdf"
    assert pdf_resource["htmlPath"] is None
    assert pdf_resource["core"]["origin"] == "INTERNAL_FILE"
    assert pdf_resource["core"]["contentAvailable"] is True

    assert link_resource["modulePath"] == "Tema 1"
    assert link_resource["itemPath"] == "Tema 1 > Enlace externo"
    assert link_resource["sourceUrl"] == "https://example.com/broken-link"
    assert link_resource["filePath"] is None
    assert link_resource["urlStatus"] == "404"
    assert link_resource["finalUrl"] == "https://example.com/broken-link"
    assert link_resource["canAccess"] is False
    assert link_resource["accessStatus"] == "NO_ACCEDE"
    assert link_resource["httpStatus"] == 404
    assert link_resource["canDownload"] is False
    assert link_resource["reasonCode"] == "NOT_FOUND"
    assert link_resource["reasonDetail"] == "La URL devolvió 404."
    assert link_resource["errorMessage"] == "La URL devolvió 404."
    assert link_resource["origin"] == "EXTERNAL_URL"
    assert link_resource["contentAvailable"] is False
    assert link_resource["localPath"] is None
    assert link_resource["htmlPath"] is None
    assert link_resource["core"]["origin"] == "EXTERNAL_URL"
    assert "broken_link" in link_resource["notes"]
    assert resources_payload["noAccessCount"] == 1
    assert resources_payload["noAccessByReason"] == {"NOT_FOUND": 1}

    broken_only_response = client.get(f"/api/jobs/{job_id}/resources?onlyBroken=true")
    assert broken_only_response.status_code == 200, broken_only_response.text
    broken_resources = broken_only_response.json()["resources"]
    assert [resource["title"] for resource in broken_resources] == ["Enlace externo"]

    access_summary_response = client.get(f"/api/jobs/{job_id}/access-summary")
    assert access_summary_response.status_code == 200, access_summary_response.text
    access_summary = access_summary_response.json()
    assert access_summary["jobId"] == job_id
    assert access_summary["status"] == "done"
    assert access_summary["progress"] == 100
    assert access_summary["total"] == 2
    assert access_summary["accessible"] == 1
    assert access_summary["downloadable"] == 1
    assert access_summary["byStatus"] == {
        "OK": 1,
        "NO_ACCEDE": 1,
        "REQUIERE_INTERACCION": 0,
        "REQUIERE_SSO": 0,
        "NO_ANALIZABLE": 0,
    }
    assert access_summary["noAccessCount"] == 1
    assert access_summary["noAccessByReason"] == {"NOT_FOUND": 1}
    assert access_summary["groups"][0]["modulePath"] == "Tema 1"

    access_response = client.get(f"/api/jobs/{job_id}/access")
    assert access_response.status_code == 200, access_response.text
    access_payload = access_response.json()
    assert access_payload["jobId"] == job_id
    assert access_payload["phase"] == "DONE"
    assert access_payload["summary"]["accessible"] == 1
    assert access_payload["summary"]["downloadable"] == 1
    assert access_payload["modules"][0]["modulePath"] == "Tema 1"
    assert [resource["title"] for resource in access_payload["modules"][0]["resources"]] == [
        "Guía",
        "Enlace externo",
    ]

    download_pdf_response = client.get(f"/api/jobs/{job_id}/resources/res-pdf/download")
    assert download_pdf_response.status_code == 200, download_pdf_response.text
    assert download_pdf_response.headers["content-type"] == "application/pdf"
    assert download_pdf_response.content.startswith(b"%PDF-1.4")

    download_external_response = client.get(f"/api/jobs/{job_id}/resources/res-link/download")
    assert download_external_response.status_code == 409, download_external_response.text
    assert download_external_response.json()["code"] == "resource_not_downloadable"


def test_offline_inventory_moves_lti_and_sso_noise_out_of_main_listing(client, monkeypatch) -> None:
    class StubURLChecker:
        def check_url(self, url: str):
            checked_at = datetime(2026, 5, 1, 8, 30, tzinfo=timezone.utc)
            return UrlCheckResult(
                url=url,
                checked=True,
                broken_link=False,
                reason="sso_redirect",
                status_code=302,
                url_status="302",
                final_url=url,
                redirect_location="https://id-provider.uoc.edu/sso/login",
                checked_at=checked_at,
            )

    monkeypatch.setattr("app.services.jobs.build_url_checker", lambda settings: StubURLChecker())

    create_response = client.post(
        "/api/jobs",
        files={"file": ("course.imscc", build_imscc_with_lti_noise(), "application/octet-stream")},
    )

    assert create_response.status_code == 201, create_response.text
    job_id = create_response.json()["jobId"]

    resources_response = client.get(f"/api/jobs/{job_id}/resources")
    assert resources_response.status_code == 200, resources_response.text
    resources_payload = resources_response.json()
    assert [resource["title"] for resource in resources_payload["resources"]] == ["Guía principal"]
    assert resources_payload["totalAnalizables"] == 1
    assert resources_payload["noAnalizablesExternos"] == 1
    assert resources_payload["tecnicosIgnorados"] >= 3

    external_block = resources_payload["nonAnalyzableExternalResources"]
    assert len(external_block) == 1
    assert external_block[0]["analysisCategory"] == "NON_ANALYZABLE_EXTERNAL"
    assert external_block[0]["title"] == "Learning Tools"
    assert external_block[0]["url"] == "https://ralti.uoc.edu/launch?tool=abc"
    assert external_block[0]["accessStatus"] == "REQUIERE_SSO"

    content_check_response = client.get(
        f"/api/jobs/{job_id}/resources/{external_block[0]['id']}/content-check"
    )
    assert content_check_response.status_code == 200, content_check_response.text
    content_check = content_check_response.json()
    assert content_check["ok"] is False
    assert content_check["resourceId"] == external_block[0]["id"]
    assert content_check["title"] == "Learning Tools"
    assert content_check["origin"] == "RALTI"
    assert content_check["contentKind"] == "NOT_ANALYZABLE"
    assert content_check["contentAvailable"] is False
    assert content_check["downloadable"] is False
    assert content_check["errorCode"] == "REQUIERE_SSO"

    access_summary_response = client.get(f"/api/jobs/{job_id}/access-summary")
    assert access_summary_response.status_code == 200, access_summary_response.text
    access_summary = access_summary_response.json()
    assert access_summary["total"] == 1
    assert access_summary["totalAnalizables"] == 1
    assert access_summary["noAnalizablesExternos"] == 1
    assert access_summary["tecnicosIgnorados"] >= 3

    access_response = client.get(f"/api/jobs/{job_id}/access")
    assert access_response.status_code == 200, access_response.text
    access_payload = access_response.json()
    assert [resource["title"] for resource in access_payload["modules"][0]["resources"]] == ["Guía principal"]
    assert len(access_payload["nonAnalyzableExternalResources"]) == 1

    summary_response = client.get(f"/api/jobs/{job_id}/summary")
    assert summary_response.status_code == 200, summary_response.text
    summary_payload = summary_response.json()
    assert summary_payload["totalResources"] == 1
    assert summary_payload["totalAnalizables"] == 1
    assert summary_payload["noAnalizablesExternos"] == 1
    assert summary_payload["tecnicosIgnorados"] >= 3


def test_offline_inventory_returns_unmapped_resources_without_technical_grouping(client) -> None:
    create_response = client.post(
        "/api/jobs",
        files={"file": ("course.imscc", build_imscc_with_unmapped_resource(), "application/octet-stream")},
    )

    assert create_response.status_code == 201, create_response.text
    job_id = create_response.json()["jobId"]

    resources_response = client.get(f"/api/jobs/{job_id}/resources")
    assert resources_response.status_code == 200, resources_response.text
    payload = resources_response.json()
    resources = payload["resources"]

    assert payload["structure"]["organizations"][0]["children"][0]["title"] == "Bloque 1"
    assert payload["structure"]["unplacedResourceIds"] == ["res-orphan"]

    mapped_resource = next(resource for resource in resources if resource["id"] == "res-guide")
    unmapped_resource = next(resource for resource in resources if resource["id"] == "res-orphan")

    assert mapped_resource["modulePath"] == "Bloque 1"
    assert mapped_resource["itemPath"] == "Bloque 1 > Guía principal"
    assert mapped_resource["sectionType"] == "structured"
    assert unmapped_resource["modulePath"] is None
    assert unmapped_resource["coursePath"] is None
    assert unmapped_resource["itemPath"] is None
    assert unmapped_resource["origin"] == "INTERNAL_PAGE"
    assert unmapped_resource["sectionType"] == "global_unplaced"
    assert unmapped_resource["contentAvailable"] is True
    assert payload["globalUnplacedCount"] == 1
    assert payload["noAccessCount"] == 0
    assert payload["noAccessByReason"] == {}


def test_offline_deep_scan_discovers_nested_local_and_external_resources(client, monkeypatch) -> None:
    class StubURLChecker:
        def check_url(self, url: str):
            checked_at = datetime(2026, 4, 30, 9, 0, tzinfo=timezone.utc)
            return UrlCheckResult(
                url=url,
                checked=True,
                broken_link=False,
                status_code=200,
                url_status="200",
                final_url=url,
                checked_at=checked_at,
                content_type="text/html; charset=utf-8",
            )

    monkeypatch.setattr("app.services.jobs.build_url_checker", lambda settings: StubURLChecker())

    create_response = client.post(
        "/api/jobs",
        files={"file": ("course.imscc", build_imscc_with_offline_deep_scan(), "application/octet-stream")},
    )

    assert create_response.status_code == 201, create_response.text
    job_id = create_response.json()["jobId"]

    resources_response = client.get(f"/api/jobs/{job_id}/resources")
    assert resources_response.status_code == 200, resources_response.text
    resources = resources_response.json()["resources"]
    assert len(resources) == 6

    by_source = {
        resource.get("filePath") or resource.get("sourceUrl") or resource["id"]: resource for resource in resources
    }

    worksheet = by_source["course/downloads/worksheet.docx"]
    slides = by_source["course/downloads/slides.pptx"]
    image = by_source["course/images/chart.png"]
    web = by_source["https://example.com/resource"]
    video = by_source["https://www.youtube.com/watch?v=demo123"]
    html_page = next(resource for resource in resources if resource["id"] == "res-html")

    assert worksheet["type"] == "DOCX"
    assert worksheet["canAccess"] is True
    assert worksheet["accessStatus"] == "OK"
    assert worksheet["httpStatus"] == 200
    assert worksheet["accessStatusCode"] == 200
    assert worksheet["canDownload"] is True
    assert worksheet["parentResourceId"] == "res-html"
    assert worksheet["parentId"] == "res-html"
    assert worksheet["modulePath"] == "Unidad 1"
    assert worksheet["moduleTitle"] == "Unidad 1"
    assert worksheet["sectionTitle"] == "Unidad 1"
    assert worksheet["origin"] == "INTERNAL_FILE"
    assert worksheet["contentAvailable"] is True
    assert worksheet["localPath"] == "course/downloads/worksheet.docx"
    assert worksheet["htmlPath"] is None
    assert worksheet["core"]["type"] == "DOCX"
    assert worksheet["core"]["origin"] == "INTERNAL_FILE"

    assert slides["canAccess"] is True
    assert slides["canDownload"] is True
    assert slides["modulePath"] == "Unidad 1"
    assert slides["moduleTitle"] == "Unidad 1"
    assert slides["sectionTitle"] == "Unidad 1"
    assert slides["origin"] == "INTERNAL_FILE"
    assert image["type"] == "IMAGE"
    assert image["canAccess"] is True
    assert image["canDownload"] is True
    assert image["origin"] == "INTERNAL_FILE"
    assert web["type"] == "WEB"
    assert web["canAccess"] is True
    assert web["canDownload"] is False
    assert web["origin"] == "EXTERNAL_URL"
    assert web["contentAvailable"] is False
    assert web["localPath"] is None
    assert web["htmlPath"] is None
    assert video["type"] == "VIDEO"
    assert video["canAccess"] is True
    assert video["canDownload"] is False
    assert video["origin"] == "EXTERNAL_URL"
    assert video["contentAvailable"] is False
    assert video["localPath"] is None
    assert video["htmlPath"] is None
    assert html_page["origin"] == "INTERNAL_PAGE"
    assert html_page["contentAvailable"] is True
    assert html_page["htmlPath"] == "course/module/page.html"
    assert html_page["discoveredChildrenCount"] == 4

    summary_response = client.get(f"/api/jobs/{job_id}/summary")
    assert summary_response.status_code == 200, summary_response.text
    summary_payload = summary_response.json()
    assert summary_payload["totalResources"] == 6
    assert summary_payload["accessibleResources"] == 6
    assert summary_payload["downloadableResources"] == 4

    access_summary_response = client.get(f"/api/jobs/{job_id}/access-summary")
    assert access_summary_response.status_code == 200, access_summary_response.text
    access_summary = access_summary_response.json()
    assert access_summary["accessible"] == 6
    assert access_summary["downloadable"] == 4

    accessibility_response = client.get(f"/api/jobs/{job_id}/accessibility")
    assert accessibility_response.status_code == 200, accessibility_response.text
    accessibility = accessibility_response.json()
    assert accessibility["summary"]["docxResourcesTotal"] == 1
    assert accessibility["summary"]["docxResourcesAnalyzed"] == 1
    assert accessibility["summary"]["videoResourcesTotal"] == 1
    assert accessibility["summary"]["videoResourcesAnalyzed"] == 1
    assert accessibility["summary"]["byType"]["DOCX"]["resourcesAnalyzed"] == 1
    assert accessibility["summary"]["byType"]["VIDEO"]["resourcesAnalyzed"] == 1
    assert any(resource["analysisType"] == "DOCX" for resource in accessibility["resources"])
    assert any(resource["analysisType"] == "VIDEO" for resource in accessibility["resources"])

    with Session(client.app.state.engine) as session:
        events = session.exec(select(JobEvent).where(JobEvent.job_id == job_id)).all()
    assert any(event.message == "Procesando accesibilidad de los documentos Word" for event in events)
    assert any(event.message == "Procesando accesibilidad de los recursos de vídeo" for event in events)


def test_offline_deep_scan_keeps_discovered_links_inside_existing_pec_section(client, monkeypatch) -> None:
    class StubURLChecker:
        def check_url(self, url: str):
            checked_at = datetime(2026, 5, 2, 10, 0, tzinfo=timezone.utc)
            return UrlCheckResult(
                url=url,
                checked=True,
                broken_link=True,
                reason="404_not_found",
                status_code=404,
                url_status="404",
                final_url=url,
                checked_at=checked_at,
            )

    monkeypatch.setattr("app.services.jobs.build_url_checker", lambda settings: StubURLChecker())

    create_response = client.post(
        "/api/jobs",
        files={"file": ("course.imscc", build_imscc_with_pec_section_deep_scan(), "application/octet-stream")},
    )

    assert create_response.status_code == 201, create_response.text
    job_id = create_response.json()["jobId"]

    resources_response = client.get(f"/api/jobs/{job_id}/resources")
    assert resources_response.status_code == 200, resources_response.text
    payload = resources_response.json()
    resources = payload["resources"]
    by_title = {resource["title"]: resource for resource in resources}

    assert payload["structure"]["organizations"][0]["title"] == "Curso PEC"
    top_sections = payload["structure"]["organizations"][0]["children"]
    assert len(top_sections) == 1
    assert top_sections[0]["title"] == "PEC 1: Actividad inicial"
    grouped_ids = collect_resource_ids(top_sections[0])
    assert "res-html" in grouped_ids
    assert by_title["Brief de la PEC"]["id"] in grouped_ids
    assert by_title["¿Cómo citar la IA…"]["id"] in grouped_ids
    assert payload["structure"]["unplacedResourceIds"] == []

    assert by_title["Brief de la PEC"]["discovered"] is True
    assert by_title["Brief de la PEC"]["parentResourceId"] == "res-html"
    assert by_title["Brief de la PEC"]["parentId"] == "res-html"
    assert by_title["Brief de la PEC"]["origin"] == "INTERNAL_FILE"
    assert by_title["Brief de la PEC"]["localPath"] == "course/downloads/brief.pdf"
    assert by_title["¿Cómo citar la IA…"]["discovered"] is True
    assert by_title["¿Cómo citar la IA…"]["parentResourceId"] == "res-html"
    assert by_title["¿Cómo citar la IA…"]["parentId"] == "res-html"
    assert by_title["¿Cómo citar la IA…"]["origin"] == "EXTERNAL_URL"
    assert by_title["¿Cómo citar la IA…"]["localPath"] is None
    assert by_title["¿Cómo citar la IA…"]["reasonCode"] == "NOT_FOUND"
    assert by_title["¿Cómo citar la IA…"]["reasonDetail"] == "La URL devolvió 404."
    assert payload["noAccessCount"] == 1
    assert payload["noAccessByReason"] == {"NOT_FOUND": 1}

    access_response = client.get(f"/api/jobs/{job_id}/access")
    assert access_response.status_code == 200, access_response.text
    access_payload = access_response.json()
    assert len(access_payload["modules"]) == 1
    assert access_payload["modules"][0]["modulePath"] == "PEC 1: Actividad inicial"
    assert access_payload["summary"]["noAccessCount"] == 1
    assert access_payload["summary"]["noAccessByReason"] == {"NOT_FOUND": 1}


def test_offline_get_resource_content_returns_html_and_binary_path(client, test_settings) -> None:
    create_response = client.post(
        "/api/jobs",
        files={"file": ("course.imscc", build_imscc_with_pec_section_deep_scan(), "application/octet-stream")},
    )

    assert create_response.status_code == 201, create_response.text
    job_id = create_response.json()["jobId"]

    resources_response = client.get(f"/api/jobs/{job_id}/resources")
    assert resources_response.status_code == 200, resources_response.text
    resources = resources_response.json()["resources"]
    html_resource = next(resource for resource in resources if resource["id"] == "res-html")
    brief_resource = next(resource for resource in resources if resource["title"] == "Brief de la PEC")

    assert html_resource["origin"] == "INTERNAL_PAGE"
    assert html_resource["htmlPath"] == "course/pec1/activity.html"
    assert html_resource["contentAvailable"] is True

    html_content = get_resource_content(job_id, "res-html", settings=test_settings)
    assert html_content.ok is True
    assert html_content.contentKind == "HTML"
    assert html_content.mimeType == "text/html"
    assert html_content.htmlContent is not None
    assert "Brief de la PEC" in html_content.htmlContent
    assert "Ficha Word" in html_content.htmlContent
    assert html_content.binaryPath is not None
    assert html_content.binaryPath.endswith("course/pec1/activity.html")

    pdf_content = get_resource_content(job_id, brief_resource["id"], settings=test_settings)
    assert pdf_content.ok is True
    assert pdf_content.contentKind == "PDF"
    assert pdf_content.mimeType == "application/pdf"
    assert pdf_content.htmlContent is None
    assert pdf_content.textContent is None
    assert pdf_content.binaryPath is not None
    assert pdf_content.binaryPath.endswith("course/downloads/brief.pdf")

    html_check_response = client.get(f"/api/jobs/{job_id}/resources/res-html/content-check")
    assert html_check_response.status_code == 200, html_check_response.text
    html_check = html_check_response.json()
    assert html_check == {
        "ok": True,
        "resourceId": "res-html",
        "title": html_resource["title"],
        "type": "WEB",
        "origin": "INTERNAL_PAGE",
        "contentKind": "HTML",
        "contentAvailable": True,
        "downloadable": False,
        "mimeType": "text/html",
        "filename": "activity.html",
        "errorCode": None,
        "errorDetail": None,
    }

    pdf_check_response = client.get(f"/api/jobs/{job_id}/resources/{brief_resource['id']}/content-check")
    assert pdf_check_response.status_code == 200, pdf_check_response.text
    pdf_check = pdf_check_response.json()
    assert pdf_check["ok"] is True
    assert pdf_check["resourceId"] == brief_resource["id"]
    assert pdf_check["title"] == "Brief de la PEC"
    assert pdf_check["type"] == "PDF"
    assert pdf_check["origin"] == "INTERNAL_FILE"
    assert pdf_check["contentKind"] == "PDF"
    assert pdf_check["contentAvailable"] is True
    assert pdf_check["downloadable"] is True
    assert pdf_check["mimeType"] == "application/pdf"
    assert pdf_check["filename"] == "brief.pdf"
    assert pdf_check["errorCode"] is None

    accessibility_response = client.get(f"/api/jobs/{job_id}/accessibility")
    assert accessibility_response.status_code == 200, accessibility_response.text
    accessibility = accessibility_response.json()
    assert accessibility["jobId"] == job_id
    assert accessibility["generatedAt"] is not None
    assert accessibility["summary"]["htmlResourcesTotal"] == 1
    assert accessibility["summary"]["htmlResourcesAnalyzed"] == 1
    assert accessibility["summary"]["pdfResourcesTotal"] == 1
    assert accessibility["summary"]["pdfResourcesAnalyzed"] == 1
    assert accessibility["summary"]["docxResourcesTotal"] == 1
    assert accessibility["summary"]["docxResourcesAnalyzed"] == 1
    assert accessibility["summary"]["byType"]["HTML"]["resourcesAnalyzed"] == 1
    assert accessibility["summary"]["byType"]["PDF"]["resourcesAnalyzed"] == 1
    assert accessibility["summary"]["byType"]["DOCX"]["resourcesAnalyzed"] == 1
    assert accessibility["summary"]["failCount"] >= 1
    assert len(accessibility["resources"]) == 3
    assert {resource["analysisType"] for resource in accessibility["resources"]} == {"HTML", "PDF", "DOCX"}
    assert any(resource["resourceId"] == "res-html" for resource in accessibility["resources"])
    assert accessibility["modules"][0]["resources"][0]["resourceId"] == "res-html"
    check_statuses = {
        check["checkId"]: check["status"] for check in accessibility["modules"][0]["resources"][0]["checks"]
    }
    assert check_statuses["html.lang"] == "FAIL"

    with Session(client.app.state.engine) as session:
        events = session.exec(select(JobEvent).where(JobEvent.job_id == job_id)).all()
    assert any(event.message == "Procesando accesibilidad de los recursos HTML" for event in events)
    assert any(event.message == "Procesando accesibilidad de los recursos PDF" for event in events)
    assert any(event.message == "Procesando accesibilidad de los recursos de vídeo" for event in events)

    executive_response = client.get(f"/api/jobs/{job_id}/executive-summary")
    assert executive_response.status_code == 200, executive_response.text
    executive = executive_response.json()
    assert executive["jobId"] == job_id
    assert executive["mode"] == "OFFLINE_IMSCC"
    assert isinstance(executive["accessibilityScore"], int)
    assert executive["priority"] in {"HIGH", "MEDIUM", "LOW"}
    assert executive["summary"]["resourcesAnalyzed"] == 3
    assert executive["summary"]["notScoredResources"] == 1
    assert executive["modules"][0]["resources"][0]["reportAnchorId"].startswith("resource-")


def test_report_generation_includes_access_html_pdf_word_and_video_accessibility_summaries(
    client,
    monkeypatch,
    test_settings,
) -> None:
    class StubURLChecker:
        def check_url(self, url: str):
            checked_at = datetime(2026, 5, 3, 8, 0, tzinfo=timezone.utc)
            return UrlCheckResult(
                url=url,
                checked=True,
                broken_link=False,
                status_code=200,
                url_status="200",
                final_url=url,
                checked_at=checked_at,
                content_type="text/html; charset=utf-8",
            )

    monkeypatch.setattr("app.services.jobs.build_url_checker", lambda settings: StubURLChecker())

    create_response = client.post(
        "/api/jobs",
        files={
            "file": (
                "course.imscc",
                build_imscc_with_pec_section_deep_scan(include_video=True),
                "application/octet-stream",
            )
        },
    )

    assert create_response.status_code == 201, create_response.text
    job_id = create_response.json()["jobId"]

    report_response = client.post(f"/api/jobs/{job_id}/report")
    assert report_response.status_code == 200, report_response.text
    report = report_response.json()

    assert report["meta"]["jobId"] == job_id
    assert report["mode"] == {"key": "OFFLINE_IMSCC", "label": "OFFLINE IMSCC"}
    assert report["accessSummary"]["resourcesDetected"] == 5
    assert report["accessSummary"]["resourcesAccessed"] == 5
    assert report["accessSummary"]["downloadable"] == 3
    assert report["accessSummary"]["noAccessible"] == 0
    assert report["accessSummary"]["requiresSSO"] == 0
    assert report["accessSummary"]["requiresInteraction"] == 0
    assert report["automaticAccessibilitySummary"]["htmlResourcesDetected"] == 1
    assert report["automaticAccessibilitySummary"]["htmlResourcesAnalyzed"] == 1
    assert report["automaticAccessibilitySummary"]["pdfResourcesDetected"] == 1
    assert report["automaticAccessibilitySummary"]["pdfResourcesAnalyzed"] == 1
    assert report["automaticAccessibilitySummary"]["wordResourcesDetected"] == 1
    assert report["automaticAccessibilitySummary"]["wordResourcesAnalyzed"] == 1
    assert report["automaticAccessibilitySummary"]["videoResourcesDetected"] == 1
    assert report["automaticAccessibilitySummary"]["videoResourcesAnalyzed"] == 1
    assert report["automaticAccessibilitySummary"]["failCount"] >= 1
    assert report["automaticAccessibilitySummary"]["warningCount"] >= 1
    assert 0 <= report["executiveSummary"]["score"] <= 100
    assert report["executiveSummary"]["priority"] in {"alta", "media", "baja"}
    assert report["executiveSummary"]["resourcesDetected"] == 5
    assert report["executiveSummary"]["resourcesAnalyzed"] == 4
    assert len(report["executiveSummary"]["priorityRecommendations"]) == 3
    assert report["moduleScores"]
    assert report["moduleScores"][0]["resourcesAnalyzed"] == 4
    assert report["resourceScores"]
    assert any(resource["type"] == "VIDEO" for resource in report["resourceScores"])
    assert report["htmlAccessibilitySummary"]["resourcesDetected"] == 1
    assert report["htmlAccessibilitySummary"]["resourcesAnalyzed"] == 1
    assert report["htmlAccessibilitySummary"]["failCount"] >= 1
    assert report["htmlAccessibilitySummary"]["warningCount"] >= 1
    assert report["pdfAccessibilitySummary"]["resourcesDetected"] == 1
    assert report["pdfAccessibilitySummary"]["resourcesAnalyzed"] == 1
    assert report["pdfAccessibilitySummary"]["failCount"] >= 1
    assert report["pdfAccessibilitySummary"]["warningCount"] >= 1
    assert report["wordAccessibilitySummary"]["resourcesDetected"] == 1
    assert report["wordAccessibilitySummary"]["resourcesAnalyzed"] == 1
    assert report["wordAccessibilitySummary"]["failCount"] >= 1
    assert report["wordAccessibilitySummary"]["warningCount"] >= 1
    assert report["videoAccessibilitySummary"]["resourcesDetected"] == 1
    assert report["videoAccessibilitySummary"]["resourcesAnalyzed"] == 1
    assert report["videoAccessibilitySummary"]["warningCount"] >= 1
    assert report["issueSummary"]
    assert report["keyIssues"]
    assert report["keyIssues"][0]["status"] == "FAIL"
    assert any(issue["status"] == "WARNING" for issue in report["keyIssues"])
    assert any(issue["resourceType"] == "PDF" for issue in report["keyIssues"])
    assert any(issue["resourceType"] == "WORD" for issue in report["keyIssues"])
    assert any(issue["resourceType"] == "VIDEO" for issue in report["keyIssues"])
    assert report["htmlResources"][0]["resourceId"] == "res-html"
    assert report["htmlResources"][0]["overallStatus"] == "FAIL"
    assert report["htmlResources"][0]["summarized"] is False
    assert report["pdfResources"][0]["title"] == "Brief de la PEC"
    assert report["pdfResources"][0]["overallStatus"] == "FAIL"
    assert report["pdfResources"][0]["summarized"] is False
    assert report["wordResources"][0]["title"] == "Ficha Word"
    assert report["wordResources"][0]["overallStatus"] == "FAIL"
    assert report["wordResources"][0]["summarized"] is False
    assert report["videoResources"][0]["title"] == "Vídeo de apoyo"
    assert report["videoResources"][0]["provider"] == "YouTube"
    assert report["videoResources"][0]["overallStatus"] == "WARNING"
    assert any(
        resource["reason"] == "EXTERNO_NO_ANALIZADO" for resource in report["notAutomaticallyAnalyzable"]
    )

    report_dir = test_settings.storage_root / "jobs" / job_id / "report"
    json_path = report_dir / "report.json"
    docx_path = report_dir / "report.docx"
    pdf_path = report_dir / "report.pdf"
    assert json_path.exists()
    assert docx_path.exists()
    assert pdf_path.exists()
    assert docx_path.stat().st_size > 0
    assert pdf_path.stat().st_size > 0

    stored_report = json.loads(json_path.read_text(encoding="utf-8"))
    assert stored_report["accessSummary"]["resourcesDetected"] == 5
    assert stored_report["automaticAccessibilitySummary"]["pdfResourcesAnalyzed"] == 1
    assert stored_report["automaticAccessibilitySummary"]["wordResourcesAnalyzed"] == 1
    assert stored_report["automaticAccessibilitySummary"]["videoResourcesAnalyzed"] == 1
    assert stored_report["executiveSummary"]["resourcesAnalyzed"] == 4
    assert stored_report["moduleScores"]
    assert stored_report["resourceScores"]
    assert stored_report["htmlAccessibilitySummary"]["resourcesAnalyzed"] == 1
    assert stored_report["pdfAccessibilitySummary"]["resourcesAnalyzed"] == 1
    assert stored_report["wordAccessibilitySummary"]["resourcesAnalyzed"] == 1
    assert stored_report["videoAccessibilitySummary"]["resourcesAnalyzed"] == 1
    assert stored_report["keyIssues"][0]["status"] == "FAIL"

    word_document = Document(str(docx_path))
    paragraphs = [paragraph.text for paragraph in word_document.paragraphs if paragraph.text]
    assert "Resumen ejecutivo" in paragraphs
    assert "Resumen de acceso" in paragraphs
    assert "Resumen de accesibilidad automática" in paragraphs
    assert "Puntuación por módulo" in paragraphs
    assert "Puntuación por recurso" in paragraphs
    assert "Principales incidencias" in paragraphs
    assert "Detalle técnico" in paragraphs
    assert "Detalle por recurso HTML" in paragraphs
    assert "Detalle por recurso PDF" in paragraphs
    assert "Detalle por recurso Word" in paragraphs
    assert "Detalle por recurso de vídeo" in paragraphs
    assert "Recursos no analizables automáticamente" in paragraphs
    assert any("Página de actividad" in paragraph for paragraph in paragraphs)
    assert any("Brief de la PEC" in paragraph for paragraph in paragraphs)
    assert any("Ficha Word" in paragraph for paragraph in paragraphs)
    assert any("Vídeo de apoyo" in paragraph for paragraph in paragraphs)


def test_checklist_upsert_is_idempotent(client, test_settings) -> None:
    job_id = "job-thread4-upsert"
    write_inventory(
        test_settings,
        job_id,
        [
            {
                "id": "course-home",
                "title": "Inicio",
                "type": "WEB",
                "origin": "imscc:webcontent",
                "path": "course/home.html",
                "course_path": "Inicio",
                "status": "OK",
            }
        ],
    )

    payload = {
        "responses": [
            {"itemKey": "keyboard", "value": "PASS"},
            {"itemKey": "focus", "value": "FAIL", "comment": "El foco apenas se ve."},
        ]
    }

    first_save = client.put(f"/api/jobs/{job_id}/resources/course-home/checklist", json=payload)
    assert first_save.status_code == 200, first_save.text

    second_save = client.put(f"/api/jobs/{job_id}/resources/course-home/checklist", json=payload)
    assert second_save.status_code == 200, second_save.text
    assert second_save.json()["reviewState"] == "NEEDS_FIX"
    assert second_save.json()["failCount"] == 1

    with Session(client.app.state.engine) as session:
        persisted = session.exec(
            select(ChecklistResponse).where(
                ChecklistResponse.job_id == job_id,
                ChecklistResponse.resource_id == "course-home",
            )
        ).all()

    assert len(persisted) == 2
    assert sorted(response.item_key for response in persisted) == ["focus", "keyboard"]


def test_review_state_transitions(client, test_settings) -> None:
    job_id = "job-thread4-states"
    write_inventory(
        test_settings,
        job_id,
        [
            {
                "id": "lesson-video",
                "title": "Video de apoyo",
                "type": "VIDEO",
                "origin": "external",
                "url": "https://example.edu/video",
                "course_path": "Modulo 1/Video",
                "status": "WARN",
            }
        ],
    )

    detail_response = client.get(f"/api/jobs/{job_id}/resources/lesson-video")
    assert detail_response.status_code == 200, detail_response.text
    item_keys = [item["itemKey"] for item in detail_response.json()["checklist"]["items"]]

    in_review_response = client.put(
        f"/api/jobs/{job_id}/resources/lesson-video/checklist",
        json={"responses": [{"itemKey": item_keys[0], "value": "PASS"}]},
    )
    assert in_review_response.status_code == 200, in_review_response.text
    assert in_review_response.json()["reviewState"] == "IN_REVIEW"
    assert in_review_response.json()["failCount"] == 0

    needs_fix_response = client.put(
        f"/api/jobs/{job_id}/resources/lesson-video/checklist",
        json={"responses": [{"itemKey": item_keys[1], "value": "FAIL"}]},
    )
    assert needs_fix_response.status_code == 200, needs_fix_response.text
    assert needs_fix_response.json()["reviewState"] == "NEEDS_FIX"
    assert needs_fix_response.json()["failCount"] == 1

    all_pass_response = client.put(
        f"/api/jobs/{job_id}/resources/lesson-video/checklist",
        json={"responses": [{"itemKey": item_key, "value": "PASS"} for item_key in item_keys]},
    )
    assert all_pass_response.status_code == 200, all_pass_response.text
    assert all_pass_response.json()["reviewState"] == "OK"
    assert all_pass_response.json()["failCount"] == 0

    resources_response = client.get(f"/api/jobs/{job_id}/resources")
    assert resources_response.status_code == 200, resources_response.text
    resource = resources_response.json()["resources"][0]
    assert resource["reviewState"] == "OK"
    assert resource["failCount"] == 0
    assert resources_response.json()["reviewSession"]["status"] == "COMPLETE"


def test_report_generation_requires_finished_job(client, test_settings) -> None:
    job_id = "5e6d779a-5df6-4cc7-b54f-b5fa7af915c9"
    job_dir = Path(test_settings.storage_root) / "jobs" / job_id
    job_dir.mkdir(parents=True, exist_ok=True)

    with Session(client.app.state.engine) as session:
        session.add(
            ProcessingJob(
                id=job_id,
                original_filename="course.imscc",
                stored_filename="course.imscc",
                size_bytes=1024,
                storage_dir=str(job_dir),
                status="processing",
                progress=50,
                current_step=2,
                total_steps=4,
                message="Procesando curso",
            )
        )
        session.commit()

    response = client.post(f"/api/jobs/{job_id}/report")
    assert response.status_code == 409, response.text
    assert response.json()["message"] == "Job aun en proceso."
