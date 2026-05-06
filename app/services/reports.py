from __future__ import annotations

from collections import Counter, defaultdict
from datetime import datetime
import json
from pathlib import Path, PurePosixPath
import shutil
import subprocess
from typing import Any
from uuid import uuid4
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fastapi import status
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlmodel import Session, select

from app.core.config import Settings
from app.core.errors import AppError
from app.models import Job as ProcessingJob
from app.models.entities import (
    ChecklistResponse,
    ChecklistValue,
    Job as ReviewJob,
    ReportRecord,
    Resource,
    ResourceHealthStatus,
    ResourceType,
    utcnow,
)
from app.schemas import GeneratedReportResponse, ReportDownloads, ReportFailure, ReportGroup, ResourceResponse
from app.services.catalog import SEVERITY_ORDER, get_item_severity
from app.services.docx_accessibility import ensure_docx_accessibility_report
from app.services.executive_summary import CRITICAL_FAIL_CHECKS, IMPORTANT_WARNING_CHECKS
from app.services.html_accessibility import ensure_accessibility_report
from app.services.pdf_accessibility import ensure_pdf_accessibility_report
from app.services.resource_core import normalize_resource
from app.services.review_service import ensure_job_inventory, ensure_review_rollups, get_templates_by_type, load_inventory_file
from app.services.storage import get_reports_dir
from app.services.video_accessibility import (
    VIDEO_ANALYSIS_SCOPE_NOTE,
    detect_video_provider,
    ensure_video_accessibility_report,
)

TYPE_LABELS = {
    ResourceType.WEB: "Web",
    ResourceType.PDF: "PDF",
    ResourceType.DOCX: "Word",
    ResourceType.VIDEO: "Video",
    ResourceType.NOTEBOOK: "Notebook",
    ResourceType.IMAGE: "Imagen",
    ResourceType.FILE: "Archivo",
    ResourceType.OTHER: "Otro",
}

STATUS_LABELS = {
    ResourceHealthStatus.OK: "OK",
    ResourceHealthStatus.WARN: "AVISO",
    ResourceHealthStatus.ERROR: "ERROR",
}

MEDIA_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "json": "application/json",
}

AUTO_STATUS_ORDER = {"FAIL": 0, "WARNING": 1}
AUTO_NOT_ANALYZABLE_EXPLANATION = (
    "No se analizan automáticamente porque requieren autenticación externa, interacción humana "
    "o un tipo de análisis todavía no implementado."
)
REPORT_NAVY = "002B45"
REPORT_CYAN = "00A6B2"
REPORT_BORDER = "C9D7DF"
REPORT_PRIORITY_ORDER = {"alta": 0, "media": 1, "baja": 2}
REPORT_TECHNICAL_STATUSES = {"FAIL", "WARNING", "ERROR"}


def _download_urls(job_id: str) -> dict[str, str]:
    return {
        "pdfUrl": f"/api/jobs/{job_id}/report/download?format=pdf",
        "docxUrl": f"/api/jobs/{job_id}/report/download?format=docx",
        "jsonUrl": f"/api/jobs/{job_id}/report/download?format=json",
    }


def _legacy_downloads(job_id: str) -> ReportDownloads:
    return ReportDownloads(
        pdfUrl=f"/api/reports/{job_id}/download/pdf",
        docxUrl=f"/api/reports/{job_id}/download/docx",
    )


def _canonical_paths(settings: Settings, job_id: str) -> tuple[Path, Path, Path]:
    reports_dir = get_reports_dir(settings, job_id)
    return reports_dir / "report.json", reports_dir / "report.docx", reports_dir / "report.pdf"


def _resource_sort_key(resource: dict[str, Any]) -> tuple[int, int, str]:
    return (-resource["stats"]["fails"], -resource["stats"]["pending"], resource["title"].lower())


def _issue_sort_key(issue: dict[str, Any]) -> tuple[int, str]:
    return (SEVERITY_ORDER.get(issue["severity"], 9), issue["label"].lower())


def _format_report_date(value: str | datetime) -> str:
    parsed = value if isinstance(value, datetime) else datetime.fromisoformat(value)
    return parsed.strftime("%Y-%m-%d %H:%M UTC")


def _stable_filename(job_id: str, created_at: str | datetime, extension: str) -> str:
    parsed = created_at if isinstance(created_at, datetime) else datetime.fromisoformat(created_at)
    return f"AccessibleCourse_Report_{job_id}_{parsed.strftime('%Y%m%d')}.{extension}"


def _enum_value(value: Any) -> str:
    if hasattr(value, "value"):
        return str(value.value)
    if value is None:
        return ""
    return str(value)


def _resolve_course_title(session: Session, job_id: str) -> str | None:
    processing_job = session.get(ProcessingJob, job_id)
    if (
        processing_job
        and getattr(processing_job, "size_bytes", 0) > 0
        and getattr(processing_job, "original_filename", None)
    ):
        stem = Path(processing_job.original_filename).stem.strip()
        if stem:
            return stem

    review_job = session.get(ReviewJob, job_id)
    if review_job and review_job.name:
        candidate = review_job.name.strip()
        if candidate:
            return candidate

    return None


def _resolve_mode(session: Session, job_id: str, inventory_items: list[Any]) -> dict[str, str]:
    if any(normalize_resource(item).origin == "ONLINE_CANVAS" for item in inventory_items):
        return {"key": "ONLINE_CANVAS", "label": "ONLINE Canvas"}

    processing_job = session.get(ProcessingJob, job_id)
    original_filename = getattr(processing_job, "original_filename", "") if processing_job is not None else ""
    if Path(original_filename).suffix.lower() in {".imscc", ".zip"}:
        return {"key": "OFFLINE_IMSCC", "label": "OFFLINE IMSCC"}

    return {"key": "OFFLINE_IMSCC", "label": "OFFLINE IMSCC"}


def _assert_job_ready(session: Session, job_id: str) -> None:
    processing_job = session.get(ProcessingJob, job_id)
    if processing_job is None:
        return

    if processing_job.status in {"created", "processing"}:
        raise AppError(
            code="job_not_ready",
            message="Job aun en proceso.",
            status_code=status.HTTP_409_CONFLICT,
            job_id=job_id,
        )
    if processing_job.status != "done":
        raise AppError(
            code="job_not_ready",
            message="El job no esta listo para generar el informe.",
            status_code=status.HTTP_409_CONFLICT,
            job_id=job_id,
        )


def _ensure_report_ready(session: Session, settings: Settings, job_id: str) -> None:
    _assert_job_ready(session, job_id)
    try:
        ensure_job_inventory(session, settings, job_id)
        ensure_review_rollups(session, job_id)
    except FileNotFoundError as exc:
        raise AppError(
            code="inventory_not_found",
            message="No hemos encontrado el inventario del curso para este job.",
            status_code=status.HTTP_404_NOT_FOUND,
            details={"reason": str(exc)},
            job_id=job_id,
        ) from exc
    except ValueError as exc:
        raise AppError(
            code="invalid_inventory",
            message="El inventario del curso no tiene un formato valido.",
            status_code=status.HTTP_409_CONFLICT,
            details={"reason": str(exc)},
            job_id=job_id,
        ) from exc


def _origin_label(resource: Resource) -> str:
    value = (resource.origin or "").lower()
    if resource.url or "extern" in value:
        return "externo"
    return "interno"


def _source_label(resource: Resource) -> str | None:
    return resource.path or resource.url or None


def _course_path(resource: Resource) -> str:
    if resource.course_path:
        return resource.course_path
    source = resource.path or resource.url or ""
    if source.startswith(("http://", "https://")):
        return "Enlaces externos"
    parent = PurePosixPath(source).parent.as_posix().strip(".")
    return parent or "Raiz del curso"


def _resource_response(resource: Resource) -> ResourceResponse:
    return ResourceResponse(
        id=resource.id,
        title=resource.title,
        type=TYPE_LABELS[resource.type],
        origin=_origin_label(resource),
        status=STATUS_LABELS[resource.status],
        href=_source_label(resource),
    )


def _recommendations(resources: list[dict[str, Any]]) -> list[str]:
    issue_counter = Counter()
    severity_counter = Counter()
    for resource in resources:
        for issue in resource["fails"] + resource["pending"]:
            issue_counter[issue["itemKey"]] += 1
            severity_counter[issue["severity"]] += 1

    recommendations = [
        "Prioriza primero los FAIL de severidad HIGH en los recursos con mayor uso docente.",
        "Convierte cada PENDING en una comprobacion verificable antes de publicar el curso.",
        "Agrupa la correccion por ruta del curso para reducir retrabajo entre equipos.",
    ]

    if issue_counter["captions"] or issue_counter["transcript"]:
        recommendations.append("Completa subtitulos y transcripciones en los videos antes de revisar mejoras menores.")
    if issue_counter["tagged"] or issue_counter["reading_order"] or issue_counter["ocr_scan"]:
        recommendations.append("Reexporta los PDF con etiquetado, orden de lectura y OCR revisados.")
    if issue_counter["alt_text"] or issue_counter["alt_images"] or issue_counter["alternative_text"]:
        recommendations.append("Añade alternativas textuales a imagenes, figuras y salidas visuales clave.")
    if issue_counter["keyboard"] or issue_counter["focus"] or issue_counter["player_controls"]:
        recommendations.append(
            "Revisa teclado y foco visible en los recursos interactivos con mayor frecuencia de uso."
        )
    if severity_counter["HIGH"] == 0 and len(recommendations) < 4:
        recommendations.append("Empieza por los recursos con mas FAIL acumulados para ganar impacto rapidamente.")

    return recommendations[:6]


def _item_course_path(item: Any) -> str:
    value = getattr(item, "item_path", None) or getattr(item, "course_path", None) or getattr(item, "module_title", None)
    if isinstance(value, str) and value.strip():
        return value.strip()
    return "Raiz del curso"


def _item_module_title(item: Any) -> str | None:
    for value in (
        getattr(item, "module_title", None),
        getattr(item, "section_title", None),
        getattr(item, "course_path", None),
    ):
        if isinstance(value, str) and value.strip():
            return value.strip()
    return None


def _item_type_label(item: Any) -> str:
    item_type = getattr(item, "type", None)
    if item_type in TYPE_LABELS:
        return TYPE_LABELS[item_type]
    return TYPE_LABELS.get(ResourceType(_enum_value(item_type)), "Otro") if _enum_value(item_type) else "Otro"


def _is_main_item(item: Any) -> bool:
    return getattr(item, "analysis_category", "MAIN_ANALYZABLE") == "MAIN_ANALYZABLE"


def _is_auxiliary_item(item: Any) -> bool:
    return getattr(item, "analysis_category", "MAIN_ANALYZABLE") == "NON_ANALYZABLE_EXTERNAL"


def _is_html_candidate(item: Any) -> bool:
    if not _is_main_item(item):
        return False
    core = normalize_resource(item)
    if core.type != "WEB":
        return False
    if core.origin in {"EXTERNAL_URL", "RALTI", "LTI"}:
        return False
    if core.htmlPath:
        return True
    local_path = core.localPath or ""
    if Path(local_path).suffix.lower() in {".html", ".htm", ".xhtml"}:
        return True
    return core.origin == "ONLINE_CANVAS"


def _is_pdf_candidate(item: Any) -> bool:
    if not _is_main_item(item):
        return False
    core = normalize_resource(item)
    if core.type != "PDF":
        return False
    if core.origin in {"EXTERNAL_URL", "RALTI", "LTI"}:
        return False
    if core.accessStatus in {"REQUIERE_SSO", "REQUIERE_INTERACCION", "NO_ANALIZABLE"}:
        return False
    return bool(core.contentAvailable)


def _is_docx_candidate(item: Any) -> bool:
    if not _is_main_item(item):
        return False
    core = normalize_resource(item)
    if core.type != "DOCX":
        return False
    if core.origin in {"EXTERNAL_URL", "RALTI", "LTI"}:
        return False
    if core.accessStatus in {"REQUIERE_SSO", "REQUIERE_INTERACCION", "NO_ANALIZABLE"}:
        return False
    return bool(core.contentAvailable)


def _is_video_candidate(item: Any) -> bool:
    if not _is_main_item(item):
        return False
    core = normalize_resource(item)
    if core.type != "VIDEO":
        return False
    if core.origin in {"RALTI", "LTI"}:
        return False
    if core.accessStatus in {"REQUIERE_SSO", "REQUIERE_INTERACCION", "NO_ANALIZABLE", "NO_ACCEDE"}:
        return False
    return True


def _build_access_summary_data(inventory_items: list[Any]) -> dict[str, int]:
    main_items = [item for item in inventory_items if _is_main_item(item)]
    relevant_items = [item for item in inventory_items if getattr(item, "analysis_category", "") != "TECHNICAL_IGNORED"]
    return {
        "resourcesDetected": len(main_items),
        "resourcesAccessed": sum(1 for item in main_items if bool(getattr(item, "can_access", False))),
        "downloadable": sum(1 for item in main_items if bool(getattr(item, "can_download", False))),
        "noAccessible": sum(1 for item in main_items if _enum_value(getattr(item, "access_status", None)) == "NO_ACCEDE"),
        "requiresSSO": sum(1 for item in relevant_items if _enum_value(getattr(item, "access_status", None)) == "REQUIERE_SSO"),
        "requiresInteraction": sum(
            1 for item in relevant_items if _enum_value(getattr(item, "access_status", None)) == "REQUIERE_INTERACCION"
        ),
        "globalUnplaced": sum(1 for item in main_items if getattr(item, "section_type", None) == "global_unplaced"),
        "noAnalyzableExternal": sum(1 for item in inventory_items if _is_auxiliary_item(item)),
        "technicalIgnored": sum(
            1 for item in inventory_items if getattr(item, "analysis_category", "") == "TECHNICAL_IGNORED"
        ),
    }


def _ensure_accessibility_data(
    settings: Settings,
    job_id: str,
    inventory_items: list[Any],
    *,
    canvas_client: Any | None = None,
    canvas_credentials: Any | None = None,
    course_id: str | None = None,
):
    payload = [item.model_dump(mode="python") for item in inventory_items]
    return ensure_accessibility_report(
        settings=settings,
        job_id=job_id,
        resources=payload,
        canvas_client=canvas_client,
        canvas_credentials=canvas_credentials,
        course_id=course_id,
    )


def _ensure_pdf_accessibility_data(
    settings: Settings,
    job_id: str,
    inventory_items: list[Any],
    *,
    canvas_client: Any | None = None,
    canvas_credentials: Any | None = None,
    course_id: str | None = None,
):
    payload = [item.model_dump(mode="python") for item in inventory_items]
    return ensure_pdf_accessibility_report(
        settings=settings,
        job_id=job_id,
        resources=payload,
        canvas_client=canvas_client,
        canvas_credentials=canvas_credentials,
        course_id=course_id,
    )


def _ensure_docx_accessibility_data(
    settings: Settings,
    job_id: str,
    inventory_items: list[Any],
    *,
    canvas_client: Any | None = None,
    canvas_credentials: Any | None = None,
    course_id: str | None = None,
):
    payload = [item.model_dump(mode="python") for item in inventory_items]
    return ensure_docx_accessibility_report(
        settings=settings,
        job_id=job_id,
        resources=payload,
        canvas_client=canvas_client,
        canvas_credentials=canvas_credentials,
        course_id=course_id,
    )


def _ensure_video_accessibility_data(
    settings: Settings,
    job_id: str,
    inventory_items: list[Any],
    *,
    canvas_client: Any | None = None,
    canvas_credentials: Any | None = None,
    course_id: str | None = None,
):
    payload = [item.model_dump(mode="python") for item in inventory_items]
    return ensure_video_accessibility_report(
        settings=settings,
        job_id=job_id,
        resources=payload,
        canvas_client=canvas_client,
        canvas_credentials=canvas_credentials,
        course_id=course_id,
    )


def _flatten_accessibility_resources(accessibility_report) -> list[tuple[str, Any]]:
    flattened: list[tuple[str, Any]] = []
    for module in accessibility_report.modules:
        for resource in module.resources:
            if getattr(resource, "analysisType", None) in {None, "HTML"}:
                flattened.append((module.title, resource))
    return flattened


def _flatten_pdf_accessibility_resources(pdf_accessibility_report) -> list[tuple[str, Any]]:
    flattened: list[tuple[str, Any]] = []
    for module in pdf_accessibility_report.modules:
        for resource in module.resources:
            if getattr(resource, "analysisType", None) == "PDF":
                flattened.append((module.title, resource))
    return flattened


def _flatten_docx_accessibility_resources(docx_accessibility_report) -> list[tuple[str, Any]]:
    flattened: list[tuple[str, Any]] = []
    for module in docx_accessibility_report.modules:
        for resource in module.resources:
            if getattr(resource, "analysisType", None) == "DOCX":
                flattened.append((module.title, resource))
    return flattened


def _flatten_video_accessibility_resources(video_accessibility_report) -> list[tuple[str, Any]]:
    flattened: list[tuple[str, Any]] = []
    for module in video_accessibility_report.modules:
        for resource in module.resources:
            if getattr(resource, "analysisType", None) == "VIDEO":
                flattened.append((module.title, resource))
    return flattened


def _build_html_summary_data(inventory_items: list[Any], accessibility_report) -> dict[str, int]:
    type_summary = accessibility_report.summary.byType.get("HTML")
    return {
        "resourcesDetected": sum(1 for item in inventory_items if _is_html_candidate(item)),
        "resourcesAnalyzed": accessibility_report.summary.htmlResourcesAnalyzed,
        "passCount": type_summary.passCount if type_summary else accessibility_report.summary.passCount,
        "failCount": type_summary.failCount if type_summary else accessibility_report.summary.failCount,
        "warningCount": type_summary.warningCount if type_summary else accessibility_report.summary.warningCount,
        "notApplicableCount": type_summary.notApplicableCount if type_summary else accessibility_report.summary.notApplicableCount,
        "errorCount": type_summary.errorCount if type_summary else accessibility_report.summary.errorCount,
    }


def _build_pdf_summary_data(inventory_items: list[Any], pdf_accessibility_report) -> dict[str, int]:
    type_summary = pdf_accessibility_report.summary.byType.get("PDF")
    return {
        "resourcesDetected": sum(1 for item in inventory_items if _is_pdf_candidate(item)),
        "resourcesAnalyzed": pdf_accessibility_report.summary.pdfResourcesAnalyzed,
        "passCount": type_summary.passCount if type_summary else pdf_accessibility_report.summary.passCount,
        "failCount": type_summary.failCount if type_summary else pdf_accessibility_report.summary.failCount,
        "warningCount": type_summary.warningCount if type_summary else pdf_accessibility_report.summary.warningCount,
        "notApplicableCount": type_summary.notApplicableCount
        if type_summary
        else pdf_accessibility_report.summary.notApplicableCount,
        "errorCount": type_summary.errorCount if type_summary else pdf_accessibility_report.summary.errorCount,
    }


def _build_docx_summary_data(inventory_items: list[Any], docx_accessibility_report) -> dict[str, int]:
    type_summary = docx_accessibility_report.summary.byType.get("DOCX")
    return {
        "resourcesDetected": sum(1 for item in inventory_items if _is_docx_candidate(item)),
        "resourcesAnalyzed": docx_accessibility_report.summary.docxResourcesAnalyzed,
        "passCount": type_summary.passCount if type_summary else docx_accessibility_report.summary.passCount,
        "failCount": type_summary.failCount if type_summary else docx_accessibility_report.summary.failCount,
        "warningCount": type_summary.warningCount if type_summary else docx_accessibility_report.summary.warningCount,
        "notApplicableCount": type_summary.notApplicableCount
        if type_summary
        else docx_accessibility_report.summary.notApplicableCount,
        "errorCount": type_summary.errorCount if type_summary else docx_accessibility_report.summary.errorCount,
    }


def _build_video_summary_data(inventory_items: list[Any], video_accessibility_report) -> dict[str, int]:
    type_summary = video_accessibility_report.summary.byType.get("VIDEO")
    return {
        "resourcesDetected": sum(1 for item in inventory_items if _is_video_candidate(item)),
        "resourcesAnalyzed": video_accessibility_report.summary.videoResourcesAnalyzed,
        "passCount": type_summary.passCount if type_summary else 0,
        "failCount": type_summary.failCount if type_summary else 0,
        "warningCount": type_summary.warningCount if type_summary else 0,
        "notApplicableCount": type_summary.notApplicableCount if type_summary else 0,
        "errorCount": type_summary.errorCount if type_summary else 0,
    }


def _build_automatic_summary_data(
    html_summary: dict[str, int],
    pdf_summary: dict[str, int],
    docx_summary: dict[str, int],
    video_summary: dict[str, int],
) -> dict[str, int]:
    return {
        "htmlResourcesDetected": html_summary["resourcesDetected"],
        "htmlResourcesAnalyzed": html_summary["resourcesAnalyzed"],
        "pdfResourcesDetected": pdf_summary["resourcesDetected"],
        "pdfResourcesAnalyzed": pdf_summary["resourcesAnalyzed"],
        "wordResourcesDetected": docx_summary["resourcesDetected"],
        "wordResourcesAnalyzed": docx_summary["resourcesAnalyzed"],
        "videoResourcesDetected": video_summary["resourcesDetected"],
        "videoResourcesAnalyzed": video_summary["resourcesAnalyzed"],
        "passCount": (
            html_summary["passCount"] + pdf_summary["passCount"] + docx_summary["passCount"] + video_summary["passCount"]
        ),
        "failCount": (
            html_summary["failCount"] + pdf_summary["failCount"] + docx_summary["failCount"] + video_summary["failCount"]
        ),
        "warningCount": (
            html_summary["warningCount"]
            + pdf_summary["warningCount"]
            + docx_summary["warningCount"]
            + video_summary["warningCount"]
        ),
        "notApplicableCount": (
            html_summary["notApplicableCount"]
            + pdf_summary["notApplicableCount"]
            + docx_summary["notApplicableCount"]
            + video_summary["notApplicableCount"]
        ),
        "errorCount": (
            html_summary["errorCount"] + pdf_summary["errorCount"] + docx_summary["errorCount"] + video_summary["errorCount"]
        ),
    }


def _build_key_issues(
    items_by_id: dict[str, Any],
    accessibility_report,
    pdf_accessibility_report,
    docx_accessibility_report,
    video_accessibility_report,
) -> list[dict[str, str | None]]:
    issues: list[dict[str, str | None]] = []
    for resource_type, resource_groups in (
        ("HTML", _flatten_accessibility_resources(accessibility_report)),
        ("PDF", _flatten_pdf_accessibility_resources(pdf_accessibility_report)),
        ("WORD", _flatten_docx_accessibility_resources(docx_accessibility_report)),
        ("VIDEO", _flatten_video_accessibility_resources(video_accessibility_report)),
    ):
        for module_title, resource in resource_groups:
            item = items_by_id.get(resource.resourceId)
            course_path = _item_course_path(item) if item is not None else module_title
            for check in resource.checks:
                if check.status not in {"FAIL", "WARNING"}:
                    continue
                issues.append(
                    {
                        "coursePath": course_path,
                        "moduleTitle": _item_module_title(item) if item is not None else module_title,
                        "resourceId": resource.resourceId,
                        "resourceTitle": resource.title,
                        "resourceType": resource_type,
                        "checkId": check.checkId,
                        "checkTitle": check.checkTitle,
                        "status": check.status,
                        "evidence": check.evidence,
                        "recommendation": check.recommendation,
                    }
                )
    return sorted(
        issues,
        key=lambda item: (
            AUTO_STATUS_ORDER.get(str(item["status"]), 9),
            str(item["coursePath"]).lower(),
            str(item["resourceTitle"]).lower(),
            str(item["resourceType"]).lower(),
            str(item["checkTitle"]).lower(),
        ),
    )


def _build_issue_summary(key_issues: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[tuple[str, str, str, str], dict[str, Any]] = {}
    for issue in key_issues:
        key = (
            str(issue["resourceType"]),
            str(issue["checkId"]),
            str(issue["checkTitle"]),
            str(issue["status"]),
        )
        group = grouped.setdefault(
            key,
            {
                "resourceType": issue["resourceType"],
                "checkId": issue["checkId"],
                "checkTitle": issue["checkTitle"],
                "status": issue["status"],
                "resourceCount": 0,
                "resources": [],
                "recommendation": issue["recommendation"],
            },
        )
        group["resourceCount"] += 1
        group["resources"].append(issue["resourceTitle"])

    for group in grouped.values():
        group["resources"] = sorted(set(group["resources"]))[:6]

    return sorted(
        grouped.values(),
        key=lambda item: (
            AUTO_STATUS_ORDER.get(str(item["status"]), 9),
            -int(item["resourceCount"]),
            str(item["resourceType"]).lower(),
            str(item["checkTitle"]).lower(),
        ),
    )


def _build_legacy_html_key_issues(items_by_id: dict[str, Any], accessibility_report) -> list[dict[str, str | None]]:
    issues: list[dict[str, str | None]] = []
    for module_title, resource in _flatten_accessibility_resources(accessibility_report):
        item = items_by_id.get(resource.resourceId)
        course_path = _item_course_path(item) if item is not None else module_title
        for check in resource.checks:
            if check.status not in {"FAIL", "WARNING"}:
                continue
            issues.append(
                {
                    "coursePath": course_path,
                    "moduleTitle": _item_module_title(item) if item is not None else module_title,
                    "resourceId": resource.resourceId,
                    "resourceTitle": resource.title,
                    "checkId": check.checkId,
                    "checkTitle": check.checkTitle,
                    "status": check.status,
                    "evidence": check.evidence,
                    "recommendation": check.recommendation,
                }
            )
    return sorted(
        issues,
        key=lambda item: (
            AUTO_STATUS_ORDER.get(str(item["status"]), 9),
            str(item["coursePath"]).lower(),
            str(item["resourceTitle"]).lower(),
            str(item["checkTitle"]).lower(),
        ),
    )


def _overall_automatic_status(checks: list[Any]) -> str:
    statuses = {check.status for check in checks}
    if "FAIL" in statuses:
        return "FAIL"
    if "WARNING" in statuses:
        return "WARNING"
    if "ERROR" in statuses:
        return "ERROR"
    return "PASS"


def _status_value(check: Any) -> str:
    if isinstance(check, dict):
        return str(check.get("status") or "")
    return str(getattr(check, "status", "") or "")


def _check_title_value(check: Any) -> str:
    if isinstance(check, dict):
        return str(check.get("checkTitle") or check.get("checkId") or "Check sin título")
    return str(getattr(check, "checkTitle", None) or getattr(check, "checkId", None) or "Check sin título")


def _check_id_value(check: Any) -> str:
    if isinstance(check, dict):
        return str(check.get("checkId") or "")
    return str(getattr(check, "checkId", "") or "")


def _check_counts(checks: list[Any]) -> Counter:
    return Counter(_status_value(check) for check in checks)


def _applicable_checks(checks: list[Any]) -> list[Any]:
    return [check for check in checks if _status_value(check) in {"PASS", "FAIL", "WARNING", "ERROR"}]


def _status_points(status: str) -> float:
    if status == "PASS":
        return 1.0
    if status == "WARNING":
        return 0.5
    return 0.0


def _score_from_checks(checks: list[Any]) -> int:
    applicable = _applicable_checks(checks)
    if not applicable:
        return 100
    obtained = sum(_status_points(_status_value(check)) for check in applicable)
    return round((obtained / len(applicable)) * 100)


def _priority_from_score(score: int, checks: list[Any] | None = None) -> str:
    applicable = _applicable_checks(checks or [])
    has_critical_fail = any(
        _status_value(check) in {"FAIL", "ERROR"} and _check_id_value(check) in CRITICAL_FAIL_CHECKS
        for check in applicable
    )
    has_important_warning = any(
        _status_value(check) == "WARNING" and _check_id_value(check) in IMPORTANT_WARNING_CHECKS
        for check in applicable
    )
    if score < 60 or has_critical_fail:
        return "alta"
    if score < 80 or has_important_warning:
        return "media"
    return "baja"


def _main_issue_from_checks(checks: list[Any]) -> str | None:
    for target_status in ("FAIL", "ERROR", "WARNING"):
        for check in checks:
            if _status_value(check) == target_status:
                return f"{_check_title_value(check)} ({target_status})"
    return None


def _score_fields(checks: list[dict[str, Any]]) -> dict[str, Any]:
    score = _score_from_checks(checks)
    return {
        "score": score,
        "priority": _priority_from_score(score, checks),
        "mainIssue": _main_issue_from_checks(checks),
    }


def _build_html_resource_details(items_by_id: dict[str, Any], accessibility_report) -> list[dict[str, Any]]:
    details: list[dict[str, Any]] = []
    for module_title, resource in _flatten_accessibility_resources(accessibility_report):
        item = items_by_id.get(resource.resourceId)
        overall_status = _overall_automatic_status(resource.checks)
        checks = [
            {
                "checkId": check.checkId,
                "checkTitle": check.checkTitle,
                "status": check.status,
                "evidence": check.evidence,
                "recommendation": check.recommendation,
            }
            for check in resource.checks
        ]
        details.append(
            {
                "resourceId": resource.resourceId,
                "title": resource.title,
                "coursePath": _item_course_path(item) if item is not None else module_title,
                "moduleTitle": _item_module_title(item) if item is not None else module_title,
                "accessStatus": resource.accessStatus,
                "overallStatus": overall_status,
                "summarized": all(check["status"] in {"PASS", "NOT_APPLICABLE"} for check in checks),
                **_score_fields(checks),
                "checks": checks,
            }
        )
    return sorted(details, key=lambda item: (str(item["coursePath"]).lower(), str(item["title"]).lower()))


def _build_pdf_resource_details(items_by_id: dict[str, Any], pdf_accessibility_report) -> list[dict[str, Any]]:
    details: list[dict[str, Any]] = []
    for module_title, resource in _flatten_pdf_accessibility_resources(pdf_accessibility_report):
        item = items_by_id.get(resource.resourceId)
        overall_status = _overall_automatic_status(resource.checks)
        checks = [
            {
                "checkId": check.checkId,
                "checkTitle": check.checkTitle,
                "status": check.status,
                "evidence": check.evidence,
                "recommendation": check.recommendation,
            }
            for check in resource.checks
        ]
        details.append(
            {
                "resourceId": resource.resourceId,
                "title": resource.title,
                "coursePath": _item_course_path(item) if item is not None else module_title,
                "moduleTitle": _item_module_title(item) if item is not None else module_title,
                "accessStatus": resource.accessStatus,
                "overallStatus": overall_status,
                "summarized": all(check["status"] in {"PASS", "NOT_APPLICABLE"} for check in checks),
                **_score_fields(checks),
                "checks": checks,
            }
        )
    return sorted(details, key=lambda item: (str(item["coursePath"]).lower(), str(item["title"]).lower()))


def _build_docx_resource_details(items_by_id: dict[str, Any], docx_accessibility_report) -> list[dict[str, Any]]:
    details: list[dict[str, Any]] = []
    for module_title, resource in _flatten_docx_accessibility_resources(docx_accessibility_report):
        item = items_by_id.get(resource.resourceId)
        overall_status = _overall_automatic_status(resource.checks)
        checks = [
            {
                "checkId": check.checkId,
                "checkTitle": check.checkTitle,
                "status": check.status,
                "evidence": check.evidence,
                "recommendation": check.recommendation,
            }
            for check in resource.checks
        ]
        details.append(
            {
                "resourceId": resource.resourceId,
                "title": resource.title,
                "coursePath": _item_course_path(item) if item is not None else module_title,
                "moduleTitle": _item_module_title(item) if item is not None else module_title,
                "accessStatus": resource.accessStatus,
                "overallStatus": overall_status,
                "summarized": all(check["status"] in {"PASS", "NOT_APPLICABLE"} for check in checks),
                **_score_fields(checks),
                "checks": checks,
            }
        )
    return sorted(details, key=lambda item: (str(item["coursePath"]).lower(), str(item["title"]).lower()))


def _build_video_resource_details(items_by_id: dict[str, Any], video_accessibility_report) -> list[dict[str, Any]]:
    details: list[dict[str, Any]] = []
    for module_title, resource in _flatten_video_accessibility_resources(video_accessibility_report):
        item = items_by_id.get(resource.resourceId)
        overall_status = _overall_automatic_status(resource.checks)
        provider, _host = detect_video_provider(item if item is not None else resource.model_dump(mode="python"))
        checks = [
            {
                "checkId": check.checkId,
                "checkTitle": check.checkTitle,
                "status": check.status,
                "evidence": check.evidence,
                "recommendation": check.recommendation,
            }
            for check in resource.checks
        ]
        details.append(
            {
                "resourceId": resource.resourceId,
                "title": resource.title,
                "coursePath": _item_course_path(item) if item is not None else module_title,
                "moduleTitle": _item_module_title(item) if item is not None else module_title,
                "accessStatus": resource.accessStatus,
                "overallStatus": overall_status,
                "provider": provider or "No identificado",
                "summarized": all(check["status"] in {"PASS", "NOT_APPLICABLE"} for check in checks),
                **_score_fields(checks),
                "checks": checks,
            }
        )
    return sorted(details, key=lambda item: (str(item["coursePath"]).lower(), str(item["title"]).lower()))


def _build_skipped_resources(inventory_items: list[Any], analyzed_ids: set[str]) -> list[dict[str, str | None]]:
    skipped: list[dict[str, str | None]] = []
    seen_ids: set[str] = set()
    for item in inventory_items:
        item_id = str(getattr(item, "id"))
        if item_id in seen_ids:
            continue

        core = normalize_resource(item)
        reason: str | None = None
        if _is_auxiliary_item(item) or core.origin in {"RALTI", "LTI"} or core.accessStatus == "REQUIERE_SSO":
            reason = "REQUIERE_SSO"
        elif core.accessStatus == "REQUIERE_INTERACCION":
            reason = "REQUIERE_INTERACCION"
        elif core.origin == "EXTERNAL_URL" and item_id not in analyzed_ids:
            reason = "EXTERNO_NO_ANALIZADO"
        elif _is_html_candidate(item) and item_id not in analyzed_ids:
            reason = "HTML_NO_ANALIZADO"
        elif _is_pdf_candidate(item) and item_id not in analyzed_ids:
            reason = "PDF_NO_ANALIZADO"
        elif _is_docx_candidate(item) and item_id not in analyzed_ids:
            reason = "WORD_NO_ANALIZADO"
        elif _is_video_candidate(item) and item_id not in analyzed_ids:
            reason = "VIDEO_NO_ANALIZADO"
        elif _is_main_item(item) and core.type not in {"WEB", "PDF", "DOCX", "VIDEO"}:
            reason = "TIPO_NO_CUBIERTO_PENDIENTE"

        if reason is None:
            continue

        seen_ids.add(item_id)
        skipped.append(
            {
                "resourceId": item_id,
                "title": str(getattr(item, "title", "Recurso sin titulo")),
                "coursePath": _item_course_path(item),
                "moduleTitle": _item_module_title(item),
                "type": _item_type_label(item),
                "origin": core.origin,
                "accessStatus": core.accessStatus,
                "reason": reason,
                "explanation": AUTO_NOT_ANALYZABLE_EXPLANATION,
            }
        )
    return sorted(skipped, key=lambda item: (str(item["coursePath"]).lower(), str(item["title"]).lower()))


def _resource_score_rows(
    html_resources: list[dict[str, Any]],
    pdf_resources: list[dict[str, Any]],
    word_resources: list[dict[str, Any]],
    video_resources: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for resource_type, resources in (
        ("HTML", html_resources),
        ("PDF", pdf_resources),
        ("WORD", word_resources),
        ("VIDEO", video_resources),
    ):
        for resource in resources:
            checks = resource.get("checks", [])
            counts = _check_counts(checks)
            rows.append(
                {
                    "resourceId": resource["resourceId"],
                    "title": resource["title"],
                    "type": resource_type,
                    "typeLabel": _auto_resource_type_label(resource_type),
                    "moduleTitle": resource.get("moduleTitle") or resource.get("coursePath") or "Raíz del curso",
                    "coursePath": resource.get("coursePath") or resource.get("moduleTitle") or "Raíz del curso",
                    "score": int(resource.get("score", _score_from_checks(checks))),
                    "priority": str(resource.get("priority") or _priority_from_score(_score_from_checks(checks), checks)),
                    "mainIssue": resource.get("mainIssue") or "Sin incidencias FAIL/WARNING",
                    "failCount": counts["FAIL"],
                    "warningCount": counts["WARNING"],
                    "errorCount": counts["ERROR"],
                }
            )

    return sorted(
        rows,
        key=lambda item: (
            REPORT_PRIORITY_ORDER.get(str(item["priority"]), 9),
            int(item["score"]),
            str(item["moduleTitle"]).lower(),
            str(item["title"]).lower(),
        ),
    )


def _module_score_rows(resource_scores: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in resource_scores:
        grouped[str(row["moduleTitle"])].append(row)

    module_rows: list[dict[str, Any]] = []
    for module_title, rows in grouped.items():
        score = round(sum(int(row["score"]) for row in rows) / len(rows)) if rows else 0
        priority = _aggregate_priority(score, rows)
        issues = [
            str(row["mainIssue"])
            for row in sorted(rows, key=lambda item: (int(item["score"]), str(item["title"]).lower()))
            if row.get("mainIssue") and row["mainIssue"] != "Sin incidencias FAIL/WARNING"
        ]
        module_rows.append(
            {
                "moduleTitle": module_title,
                "score": score,
                "priority": priority,
                "resourcesAnalyzed": len(rows),
                "mainIssues": list(dict.fromkeys(issues))[:3] or ["Sin incidencias principales"],
            }
        )

    return sorted(
        module_rows,
        key=lambda item: (
            REPORT_PRIORITY_ORDER.get(str(item["priority"]), 9),
            int(item["score"]),
            str(item["moduleTitle"]).lower(),
        ),
    )


def _aggregate_priority(score: int, resource_scores: list[dict[str, Any]]) -> str:
    priorities = {str(row.get("priority")) for row in resource_scores}
    if score < 60 or "alta" in priorities:
        return "alta"
    if score < 80 or "media" in priorities:
        return "media"
    return "baja"


def _global_score(resource_scores: list[dict[str, Any]], access_summary: dict[str, int]) -> int:
    detected = int(access_summary.get("resourcesDetected", 0))
    if not resource_scores:
        return 0 if detected else 100
    return round(sum(int(row["score"]) for row in resource_scores) / len(resource_scores))


def _build_main_problem_labels(issue_summary: list[dict[str, Any]]) -> list[str]:
    labels = []
    for issue in issue_summary[:5]:
        labels.append(
            f"{_auto_resource_type_label(issue['resourceType'])}: {issue['checkTitle']} "
            f"({issue['status']}, {issue['resourceCount']} recurso(s))"
        )
    return labels or ["No se han detectado incidencias FAIL/WARNING en el análisis automático."]


def _executive_recommendations(issue_summary: list[dict[str, Any]], skipped_count: int) -> list[str]:
    recommendations = []
    for issue in issue_summary:
        recommendation = str(issue.get("recommendation") or "").strip()
        if recommendation and recommendation not in recommendations:
            recommendations.append(recommendation)
        if len(recommendations) == 3:
            return recommendations

    defaults = [
        "Corregir primero los checks FAIL que afectan a más recursos o a módulos completos.",
        "Revisar manualmente los recursos no analizables y documentar la evidencia obtenida.",
        "Regenerar el informe después de aplicar cambios para verificar la mejora de score.",
    ]
    if skipped_count == 0:
        defaults[1] = "Mantener una revisión manual breve para validar calidad pedagógica y experiencia de uso."
    for recommendation in defaults:
        if recommendation not in recommendations:
            recommendations.append(recommendation)
        if len(recommendations) == 3:
            break
    return recommendations[:3]


def _executive_narrative(main_problems: list[str]) -> str:
    if not main_problems or main_problems[0].startswith("No se han detectado"):
        return "No se han detectado barreras automáticas críticas; conviene completar la revisión manual y mantener evidencias."
    compact = "; ".join(main_problems[:3])
    return f"Las principales barreras detectadas están relacionadas con {compact}."


def _build_executive_summary(
    access_summary: dict[str, int],
    resource_scores: list[dict[str, Any]],
    issue_summary: list[dict[str, Any]],
    skipped_resources: list[dict[str, Any]],
) -> dict[str, Any]:
    score = _global_score(resource_scores, access_summary)
    main_problems = _build_main_problem_labels(issue_summary)
    return {
        "score": score,
        "priority": _aggregate_priority(score, resource_scores),
        "resourcesDetected": int(access_summary.get("resourcesDetected", 0)),
        "resourcesAnalyzed": len(resource_scores),
        "notAutomaticallyAnalyzable": len(skipped_resources),
        "mainProblems": main_problems,
        "priorityRecommendations": _executive_recommendations(issue_summary, len(skipped_resources)),
        "narrative": _executive_narrative(main_problems),
    }


def _build_manual_review_sections(
    session: Session,
    job_id: str,
    *,
    include_pending: bool,
    only_fails: bool,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], int, int]:
    template_map = get_templates_by_type(session)
    resources = session.exec(select(Resource).where(Resource.job_id == job_id).order_by(Resource.title)).all()
    responses = session.exec(select(ChecklistResponse).where(ChecklistResponse.job_id == job_id)).all()
    responses_by_resource: dict[str, list[ChecklistResponse]] = defaultdict(list)
    for response in responses:
        responses_by_resource[response.resource_id].append(response)

    resource_sections: list[dict[str, Any]] = []
    total_fails = 0
    total_pending = 0

    for resource in resources:
        template_bundle = template_map.get(resource.type) or template_map.get(ResourceType.OTHER)
        if template_bundle is None:
            continue

        responses_by_key = {response.item_key: response for response in responses_by_resource.get(resource.id, [])}
        fails: list[dict[str, Any]] = []
        pending: list[dict[str, Any]] = []

        for template_item in template_bundle.items:
            response = responses_by_key.get(template_item.key)
            value = response.value if response is not None else ChecklistValue.PENDING
            issue = {
                "itemKey": template_item.key,
                "label": template_item.label,
                "description": template_item.description or template_item.label,
                "recommendation": template_item.recommendation,
                "severity": get_item_severity(template_item.key),
                "comment": response.comment if response is not None else None,
            }

            if value == ChecklistValue.FAIL:
                fails.append({**issue, "status": "FAIL"})
            elif value == ChecklistValue.PENDING and include_pending and not only_fails:
                pending.append({**issue, "status": "PENDING"})

        fails.sort(key=_issue_sort_key)
        pending.sort(key=_issue_sort_key)
        total_fails += len(fails)
        total_pending += len(pending)

        if not fails and not pending:
            continue

        resource_sections.append(
            {
                "resourceId": resource.id,
                "title": resource.title,
                "type": TYPE_LABELS[resource.type],
                "origin": _origin_label(resource),
                "status": STATUS_LABELS[resource.status],
                "source": _source_label(resource),
                "coursePath": _course_path(resource),
                "stats": {"resources": 1, "fails": len(fails), "pending": len(pending)},
                "fails": fails,
                "pending": pending,
            }
        )

    route_groups: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for resource in resource_sections:
        route_groups[resource["coursePath"]].append(resource)

    routes = []
    for course_path in sorted(route_groups):
        grouped_resources = sorted(route_groups[course_path], key=_resource_sort_key)
        routes.append(
            {
                "coursePath": course_path,
                "stats": {
                    "resources": len(grouped_resources),
                    "fails": sum(item["stats"]["fails"] for item in grouped_resources),
                    "pending": sum(item["stats"]["pending"] for item in grouped_resources),
                },
                "resources": grouped_resources,
            }
        )

    return routes, sorted(resource_sections, key=_resource_sort_key), total_fails, total_pending


def _build_report_payload(
    session: Session,
    settings: Settings,
    job_id: str,
    *,
    include_pending: bool = True,
    only_fails: bool = False,
    canvas_client: Any | None = None,
    canvas_credentials: Any | None = None,
    course_id: str | None = None,
) -> dict[str, Any]:
    _ensure_report_ready(session, settings, job_id)

    inventory_items = load_inventory_file(settings, job_id)
    accessibility_report = _ensure_accessibility_data(
        settings,
        job_id,
        inventory_items,
        canvas_client=canvas_client,
        canvas_credentials=canvas_credentials,
        course_id=course_id,
    )
    pdf_accessibility_report = _ensure_pdf_accessibility_data(
        settings,
        job_id,
        inventory_items,
        canvas_client=canvas_client,
        canvas_credentials=canvas_credentials,
        course_id=course_id,
    )
    docx_accessibility_report = _ensure_docx_accessibility_data(
        settings,
        job_id,
        inventory_items,
        canvas_client=canvas_client,
        canvas_credentials=canvas_credentials,
        course_id=course_id,
    )
    video_accessibility_report = _ensure_video_accessibility_data(
        settings,
        job_id,
        inventory_items,
        canvas_client=canvas_client,
        canvas_credentials=canvas_credentials,
        course_id=course_id,
    )
    items_by_id = {str(item.id): item for item in inventory_items}

    created_at = utcnow()
    created_at_iso = created_at.isoformat()
    report_id = f"report-{uuid4().hex[:12]}"
    course_title = _resolve_course_title(session, job_id)
    mode = _resolve_mode(session, job_id, inventory_items)
    access_summary = _build_access_summary_data(inventory_items)
    html_summary = _build_html_summary_data(inventory_items, accessibility_report)
    pdf_summary = _build_pdf_summary_data(inventory_items, pdf_accessibility_report)
    docx_summary = _build_docx_summary_data(inventory_items, docx_accessibility_report)
    video_summary = _build_video_summary_data(inventory_items, video_accessibility_report)
    automatic_summary = _build_automatic_summary_data(html_summary, pdf_summary, docx_summary, video_summary)
    key_issues = _build_key_issues(
        items_by_id,
        accessibility_report,
        pdf_accessibility_report,
        docx_accessibility_report,
        video_accessibility_report,
    )
    issue_summary = _build_issue_summary(key_issues)
    html_resources = _build_html_resource_details(items_by_id, accessibility_report)
    pdf_resources = _build_pdf_resource_details(items_by_id, pdf_accessibility_report)
    docx_resources = _build_docx_resource_details(items_by_id, docx_accessibility_report)
    video_resources = _build_video_resource_details(items_by_id, video_accessibility_report)
    analyzed_ids = {
        resource["resourceId"] for resource in html_resources + pdf_resources + docx_resources + video_resources
    }
    skipped_resources = _build_skipped_resources(inventory_items, analyzed_ids)
    resource_scores = _resource_score_rows(html_resources, pdf_resources, docx_resources, video_resources)
    module_scores = _module_score_rows(resource_scores)

    routes, resource_sections, total_fails, total_pending = _build_manual_review_sections(
        session,
        job_id,
        include_pending=include_pending,
        only_fails=only_fails,
    )
    top_resources = sorted(
        [
            {
                "resourceId": resource["resourceId"],
                "title": resource["title"],
                "coursePath": resource["coursePath"],
                "failCount": resource["stats"]["fails"],
            }
            for resource in resource_sections
            if resource["stats"]["fails"] > 0
        ],
        key=lambda item: (-item["failCount"], item["title"].lower()),
    )[:5]

    return {
        "reportId": report_id,
        "createdAt": created_at_iso,
        "files": _download_urls(job_id),
        "stats": {
            "resources": len(session.exec(select(Resource).where(Resource.job_id == job_id)).all()),
            "fails": total_fails,
            "pending": total_pending,
        },
        "meta": {
            "reportId": report_id,
            "createdAt": created_at_iso,
            "courseTitle": course_title,
            "jobId": job_id,
            "includePending": include_pending and not only_fails,
            "onlyFails": only_fails,
            "systemVersion": settings.version,
        },
        "mode": mode,
        "accessSummary": access_summary,
        "automaticAccessibilitySummary": automatic_summary,
        "executiveSummary": _build_executive_summary(access_summary, resource_scores, issue_summary, skipped_resources),
        "moduleScores": module_scores,
        "resourceScores": resource_scores,
        "htmlAccessibilitySummary": html_summary,
        "pdfAccessibilitySummary": pdf_summary,
        "wordAccessibilitySummary": docx_summary,
        "videoAccessibilitySummary": video_summary,
        "issueSummary": issue_summary,
        "keyIssues": key_issues,
        "htmlResources": html_resources,
        "pdfResources": pdf_resources,
        "wordResources": docx_resources,
        "videoResources": video_resources,
        "notAutomaticallyAnalyzable": skipped_resources,
        "summary": {
            "resources": len(session.exec(select(Resource).where(Resource.job_id == job_id)).all()),
            "fails": total_fails,
            "pending": total_pending,
            "topResources": top_resources,
            "recommendations": _recommendations(resource_sections),
        },
        "routes": routes,
        "resources": resource_sections,
        "appendix": {
            "statusDefinitions": {
                "PENDING": "Pendiente de revisar o sin evidencia suficiente todavia.",
                "PASS": "Cumple el criterio revisado.",
                "FAIL": "No cumple el criterio y requiere accion correctiva.",
            },
            "createdAt": created_at_iso,
            "systemVersion": settings.version,
        },
    }


def _configure_docx_styles(document: Document) -> None:
    def set_style_font(style_name: str, size: int) -> None:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        if style_name.startswith("Heading"):
            style.font.color.rgb = RGBColor(0, 43, 69)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Calibri")

    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    set_style_font("Normal", 11)
    set_style_font("Heading 1", 16)
    set_style_font("Heading 2", 13)


def _docx_set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _docx_style_header_cell(cell: Any) -> None:
    _docx_set_cell_shading(cell, REPORT_NAVY)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)


def _append_docx_table(document: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0].cells
    for index, label in enumerate(headers):
        header[index].text = str(label)
        _docx_style_header_cell(header[index])

    for values in rows:
        row = table.add_row().cells
        for index, value in enumerate(values):
            row[index].text = str(value if value is not None else "-")


def _append_docx_summary_table(document: Document, rows: list[tuple[str, str]]) -> None:
    _append_docx_table(document, ["Indicador", "Valor"], [[label, value] for label, value in rows])


def _auto_resource_type_label(value: Any) -> str:
    normalized = str(value or "").upper()
    if normalized in {"DOCX", "WORD"}:
        return "Word"
    if normalized == "VIDEO":
        return "Vídeo"
    return normalized


def _compact_detail_recommendation(check: dict[str, Any]) -> str:
    if check.get("status") in {"FAIL", "WARNING"}:
        return "Ver recomendación en el resumen agrupado."
    return str(check.get("recommendation") or "")


def _clip_text(value: Any, limit: int = 260) -> str:
    text = " ".join(str(value if value is not None else "").split())
    return text if len(text) <= limit else f"{text[: limit - 1]}…"


def _technical_checks(resource: dict[str, Any]) -> list[dict[str, Any]]:
    return [check for check in resource.get("checks", []) if check.get("status") in REPORT_TECHNICAL_STATUSES]


def _priority_label(value: Any) -> str:
    normalized = str(value or "").lower()
    return {"alta": "Alta", "media": "Media", "baja": "Baja"}.get(normalized, str(value or "-"))


def _append_docx_technical_resources(
    document: Document,
    title: str,
    resources: list[dict[str, Any]],
    *,
    intro: str | None = None,
) -> None:
    document.add_heading(title, level=1)
    if intro:
        document.add_paragraph(intro)
    if not resources:
        document.add_paragraph("No hay recursos de este tipo analizados automáticamente.")
        return

    for resource in resources:
        checks = _technical_checks(resource)
        document.add_heading(resource["title"], level=2)
        metadata = (
            f"Módulo/sección: {resource['moduleTitle'] or resource['coursePath']} | "
            f"Score: {resource.get('score', 0)}/100 | Prioridad: {_priority_label(resource.get('priority'))} | "
            f"Estado de acceso: {resource['accessStatus']}"
        )
        if resource.get("provider"):
            metadata = f"{metadata} | Proveedor: {resource['provider']}"
        document.add_paragraph(metadata)
        if not checks:
            document.add_paragraph("Sin incidencias FAIL/WARNING/ERROR. Los checks PASS y NOT_APPLICABLE se omiten para compactar.")
            continue
        _append_docx_table(
            document,
            ["Check", "Estado", "Evidencia", "Recomendación"],
            [
                [
                    check["checkTitle"],
                    check["status"],
                    _clip_text(check["evidence"]),
                    _compact_detail_recommendation(check),
                ]
                for check in checks
            ],
        )


def _write_docx(destination: Path, report: dict[str, Any], brand_name: str) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_docx_styles(document)
    executive = report["executiveSummary"]

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Informe de accesibilidad")
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(0, 43, 69)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"{brand_name} · {report['meta']['courseTitle'] or report['meta']['jobId']}")
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(0, 166, 178)
    document.add_paragraph(f"Modo: {report['mode']['label']}")
    document.add_paragraph(f"Fecha: {_format_report_date(report['createdAt'])}")
    document.add_paragraph(f"Versión AccessibleCourse: {report['meta']['systemVersion']}")
    document.add_paragraph(f"Job ID: {report['meta']['jobId']}")
    score_paragraph = document.add_paragraph()
    score_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    score_run = score_paragraph.add_run(f"{executive['score']}/100")
    score_run.bold = True
    score_run.font.size = Pt(30)
    score_run.font.color.rgb = RGBColor(0, 43, 69)
    priority_paragraph = document.add_paragraph()
    priority_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    priority_run = priority_paragraph.add_run(f"Prioridad global: {_priority_label(executive['priority'])}")
    priority_run.bold = True
    priority_run.font.color.rgb = RGBColor(0, 166, 178)
    document.add_page_break()

    document.add_heading("Resumen ejecutivo", level=1)
    document.add_paragraph(executive["narrative"])
    _append_docx_summary_table(
        document,
        [
            ("Puntuación de accesibilidad analizada", f"{executive['score']}/100"),
            ("Prioridad global", _priority_label(executive["priority"])),
            ("Recursos detectados", str(executive["resourcesDetected"])),
            ("Recursos analizados", str(executive["resourcesAnalyzed"])),
            ("Recursos no analizables", str(executive["notAutomaticallyAnalyzable"])),
        ],
    )
    document.add_heading("Principales problemas", level=2)
    for problem in executive["mainProblems"]:
        document.add_paragraph(problem, style="List Bullet")
    document.add_heading("3 recomendaciones prioritarias", level=2)
    for recommendation in executive["priorityRecommendations"]:
        document.add_paragraph(recommendation, style="List Bullet")

    document.add_heading("Resumen de acceso", level=1)
    _append_docx_summary_table(
        document,
        [
            ("Recursos detectados", str(report["accessSummary"]["resourcesDetected"])),
            ("Recursos accedidos", str(report["accessSummary"]["resourcesAccessed"])),
            ("Descargables", str(report["accessSummary"]["downloadable"])),
            ("No accesibles", str(report["accessSummary"]["noAccessible"])),
            ("Requieren SSO", str(report["accessSummary"]["requiresSSO"])),
            ("Requieren interacción", str(report["accessSummary"]["requiresInteraction"])),
            ("Globales / no ubicados", str(report["accessSummary"]["globalUnplaced"])),
        ],
    )

    document.add_heading("Resumen de accesibilidad automática", level=1)
    _append_docx_summary_table(
        document,
        [
            ("Recursos HTML detectados", str(report["automaticAccessibilitySummary"]["htmlResourcesDetected"])),
            ("Recursos HTML analizados", str(report["automaticAccessibilitySummary"]["htmlResourcesAnalyzed"])),
            ("Recursos PDF detectados", str(report["automaticAccessibilitySummary"]["pdfResourcesDetected"])),
            ("Recursos PDF analizados", str(report["automaticAccessibilitySummary"]["pdfResourcesAnalyzed"])),
            ("Recursos Word detectados", str(report["automaticAccessibilitySummary"]["wordResourcesDetected"])),
            ("Recursos Word analizados", str(report["automaticAccessibilitySummary"]["wordResourcesAnalyzed"])),
            ("Recursos de vídeo detectados", str(report["automaticAccessibilitySummary"]["videoResourcesDetected"])),
            ("Recursos de vídeo analizados", str(report["automaticAccessibilitySummary"]["videoResourcesAnalyzed"])),
            ("Total PASS", str(report["automaticAccessibilitySummary"]["passCount"])),
            ("Total FAIL", str(report["automaticAccessibilitySummary"]["failCount"])),
            ("Total WARNING", str(report["automaticAccessibilitySummary"]["warningCount"])),
            ("Total NOT_APPLICABLE", str(report["automaticAccessibilitySummary"]["notApplicableCount"])),
            ("Total ERROR", str(report["automaticAccessibilitySummary"]["errorCount"])),
        ],
    )

    document.add_heading("Puntuación por módulo", level=1)
    if report["moduleScores"]:
        _append_docx_table(
            document,
            ["Módulo/sección", "Score", "Prioridad", "Recursos analizados", "Incidencias principales"],
            [
                [
                    row["moduleTitle"],
                    f"{row['score']}/100",
                    _priority_label(row["priority"]),
                    row["resourcesAnalyzed"],
                    _clip_text("; ".join(row["mainIssues"]), 180),
                ]
                for row in report["moduleScores"]
            ],
        )
    else:
        document.add_paragraph("No hay recursos analizados automáticamente para calcular score por módulo.")

    document.add_heading("Puntuación por recurso", level=1)
    if report["resourceScores"]:
        _append_docx_table(
            document,
            ["Recurso", "Tipo", "Módulo", "Score", "Prioridad", "Incidencia principal"],
            [
                [
                    row["title"],
                    row["typeLabel"],
                    row["moduleTitle"],
                    f"{row['score']}/100",
                    _priority_label(row["priority"]),
                    _clip_text(row["mainIssue"], 160),
                ]
                for row in report["resourceScores"]
            ],
        )
    else:
        document.add_paragraph("No hay recursos analizados automáticamente para calcular score por recurso.")

    document.add_heading("Principales incidencias", level=1)
    if not report["issueSummary"]:
        document.add_paragraph("No se han detectado incidencias FAIL o WARNING en los checks automáticos.")
    else:
        _append_docx_table(
            document,
            ["Tipo", "Check", "Estado", "Recursos afectados", "Recomendación"],
            [
                [
                    _auto_resource_type_label(issue_group["resourceType"]),
                    issue_group["checkTitle"],
                    issue_group["status"],
                    issue_group["resourceCount"],
                    _clip_text(issue_group["recommendation"], 220),
                ]
                for issue_group in report["issueSummary"]
            ],
        )

    document.add_page_break()
    document.add_heading("Detalle técnico", level=1)
    document.add_paragraph("Los checks PASS se omiten en esta sección para mantener el informe accionable.")
    _append_docx_technical_resources(document, "Detalle por recurso HTML", report["htmlResources"])
    _append_docx_technical_resources(document, "Detalle por recurso PDF", report["pdfResources"])
    _append_docx_technical_resources(document, "Detalle por recurso Word", report["wordResources"])
    _append_docx_technical_resources(
        document,
        "Detalle por recurso de vídeo",
        report["videoResources"],
        intro=VIDEO_ANALYSIS_SCOPE_NOTE,
    )

    document.add_heading("Recursos no analizables automáticamente", level=1)
    if not report["notAutomaticallyAnalyzable"]:
        document.add_paragraph("No hay recursos pendientes de análisis automático por autenticación, interacción o cobertura.")
    else:
        document.add_paragraph(AUTO_NOT_ANALYZABLE_EXPLANATION)
        _append_docx_table(
            document,
            ["Recurso", "Tipo", "Motivo", "Módulo/sección"],
            [
                [
                    resource["title"],
                    resource["type"],
                    resource["reason"],
                    resource["moduleTitle"] or resource["coursePath"],
                ]
                for resource in report["notAutomaticallyAnalyzable"]
            ],
        )

    document.add_heading("Revisión manual complementaria", level=1)
    if not report["routes"]:
        document.add_paragraph("No hay hallazgos FAIL o PENDING en la checklist manual.")
    else:
        for route in report["routes"]:
            document.add_heading(f"Ruta: {route['coursePath']}", level=2)
            for resource in route["resources"]:
                document.add_paragraph(
                    f"{resource['title']} | Tipo: {resource['type']} | Origen: {resource['origin']}"
                )
                table = document.add_table(rows=1, cols=5)
                table.style = "Table Grid"
                header = table.rows[0].cells
                header[0].text = "Estado"
                header[1].text = "Severidad"
                header[2].text = "Descripción"
                header[3].text = "Cómo arreglarlo"
                header[4].text = "Notas"
                for cell in header:
                    _docx_style_header_cell(cell)
                for issue in resource["fails"] + resource["pending"]:
                    row = table.add_row().cells
                    row[0].text = issue["status"]
                    row[1].text = issue["severity"]
                    row[2].text = issue["description"]
                    row[3].text = issue["recommendation"] or "Sin recomendación disponible."
                    row[4].text = issue.get("comment") or "-"

    document.add_heading("Apéndice", level=1)
    for key, value in report["appendix"]["statusDefinitions"].items():
        document.add_paragraph(f"{key}: {value}", style="List Bullet")
    document.add_paragraph(f"Fecha: {_format_report_date(report['appendix']['createdAt'])}")
    document.add_paragraph(f"Versión del sistema: {report['appendix']['systemVersion']}")
    document.save(destination)


def _pdf_text(value: Any) -> str:
    return xml_escape(str(value if value is not None else ""))


def _pdf_paragraph(value: Any, style: ParagraphStyle) -> Paragraph:
    return Paragraph(_pdf_text(value), style)


def _append_pdf_table(story: list[Any], rows: list[list[Any]], *, widths: list[float]) -> None:
    header_style = ParagraphStyle(
        name="ReportTableHeader",
        fontName="Helvetica-Bold",
        fontSize=8.4,
        leading=10,
        textColor=colors.white,
    )
    cell_style = ParagraphStyle(name="ReportTableCell", fontSize=8.2, leading=10)
    escaped_rows = [
        [Paragraph(_pdf_text(cell), header_style if row_index == 0 else cell_style) for cell in row]
        for row_index, row in enumerate(rows)
    ]
    table = Table(escaped_rows, colWidths=widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{REPORT_NAVY}")),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor(f"#{REPORT_BORDER}")),
                ("PADDING", (0, 0), (-1, -1), 5),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.extend([table, Spacer(1, 0.25 * cm)])


def _append_pdf_bullets(story: list[Any], items: list[str], style: ParagraphStyle) -> None:
    for item in items:
        story.append(_pdf_paragraph(f"• {item}", style))


def _append_pdf_technical_resources(
    story: list[Any],
    styles: dict[str, ParagraphStyle],
    title: str,
    resources: list[dict[str, Any]],
    *,
    intro: str | None = None,
) -> None:
    story.append(_pdf_paragraph(title, styles["Section"]))
    if intro:
        story.append(_pdf_paragraph(intro, styles["Normal"]))
    if not resources:
        story.append(_pdf_paragraph("No hay recursos de este tipo analizados automáticamente.", styles["Normal"]))
        return

    for resource in resources:
        checks = _technical_checks(resource)
        story.append(_pdf_paragraph(resource["title"], styles["Heading2"]))
        metadata = (
            f"Módulo/sección: {resource['moduleTitle'] or resource['coursePath']} | "
            f"Score: {resource.get('score', 0)}/100 | Prioridad: {_priority_label(resource.get('priority'))} | "
            f"Estado de acceso: {resource['accessStatus']}"
        )
        if resource.get("provider"):
            metadata = f"{metadata} | Proveedor: {resource['provider']}"
        story.append(_pdf_paragraph(metadata, styles["Normal"]))
        if not checks:
            story.append(
                _pdf_paragraph(
                    "Sin incidencias FAIL/WARNING/ERROR. Los checks PASS y NOT_APPLICABLE se omiten para compactar.",
                    styles["Normal"],
                )
            )
            continue
        rows = [["Check", "Estado", "Evidencia", "Recomendación"]]
        for check in checks:
            rows.append(
                [
                    check["checkTitle"],
                    check["status"],
                    _clip_text(check["evidence"]),
                    _compact_detail_recommendation(check),
                ]
            )
        _append_pdf_table(story, rows, widths=[3.4 * cm, 2.4 * cm, 5.3 * cm, 4.2 * cm])


def _write_pdf(destination: Path, report: dict[str, Any], brand_name: str) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CoverTitle",
            parent=styles["Title"],
            fontSize=26,
            leading=30,
            textColor=colors.HexColor(f"#{REPORT_NAVY}"),
            alignment=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverSubtitle",
            parent=styles["Normal"],
            fontSize=12,
            leading=15,
            textColor=colors.HexColor(f"#{REPORT_CYAN}"),
            alignment=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Score",
            parent=styles["Title"],
            fontSize=28,
            leading=32,
            textColor=colors.HexColor(f"#{REPORT_NAVY}"),
            alignment=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Section",
            parent=styles["Heading1"],
            fontSize=15,
            leading=18,
            textColor=colors.HexColor(f"#{REPORT_NAVY}"),
            spaceBefore=8,
            spaceAfter=6,
        )
    )
    styles["Heading2"].textColor = colors.HexColor(f"#{REPORT_NAVY}")
    executive = report["executiveSummary"]

    story: list[Any] = [
        _pdf_paragraph("Informe de accesibilidad", styles["CoverTitle"]),
        _pdf_paragraph(f"{brand_name} · {report['meta']['courseTitle'] or report['meta']['jobId']}", styles["CoverSubtitle"]),
        Spacer(1, 0.4 * cm),
        _pdf_paragraph(f"{executive['score']}/100", styles["Score"]),
        _pdf_paragraph(f"Prioridad global: {_priority_label(executive['priority'])}", styles["CoverSubtitle"]),
        Spacer(1, 0.5 * cm),
        _pdf_paragraph(f"Modo: {report['mode']['label']}", styles["Normal"]),
        _pdf_paragraph(f"Fecha: {_format_report_date(report['createdAt'])}", styles["Normal"]),
        _pdf_paragraph(f"Versión AccessibleCourse: {report['meta']['systemVersion']}", styles["Normal"]),
        _pdf_paragraph(f"Job ID: {report['meta']['jobId']}", styles["Normal"]),
        PageBreak(),
        _pdf_paragraph("Resumen ejecutivo", styles["Section"]),
    ]

    story.append(_pdf_paragraph(executive["narrative"], styles["Normal"]))
    _append_pdf_table(
        story,
        [
            ["Indicador", "Valor"],
            ["Puntuación de accesibilidad analizada", f"{executive['score']}/100"],
            ["Prioridad global", _priority_label(executive["priority"])],
            ["Recursos detectados", str(executive["resourcesDetected"])],
            ["Recursos analizados", str(executive["resourcesAnalyzed"])],
            ["Recursos no analizables", str(executive["notAutomaticallyAnalyzable"])],
        ],
        widths=[9 * cm, 4.5 * cm],
    )
    story.append(_pdf_paragraph("Principales problemas", styles["Heading2"]))
    _append_pdf_bullets(story, executive["mainProblems"], styles["Normal"])
    story.append(_pdf_paragraph("3 recomendaciones prioritarias", styles["Heading2"]))
    _append_pdf_bullets(story, executive["priorityRecommendations"], styles["Normal"])

    story.append(_pdf_paragraph("Resumen de acceso", styles["Section"]))
    _append_pdf_table(
        story,
        [
            ["Indicador", "Valor"],
            ["Recursos detectados", str(report["accessSummary"]["resourcesDetected"])],
            ["Recursos accedidos", str(report["accessSummary"]["resourcesAccessed"])],
            ["Descargables", str(report["accessSummary"]["downloadable"])],
            ["No accesibles", str(report["accessSummary"]["noAccessible"])],
            ["Requieren SSO", str(report["accessSummary"]["requiresSSO"])],
            ["Requieren interacción", str(report["accessSummary"]["requiresInteraction"])],
            ["Globales / no ubicados", str(report["accessSummary"]["globalUnplaced"])],
        ],
        widths=[9 * cm, 4.2 * cm],
    )

    story.append(_pdf_paragraph("Resumen de accesibilidad automática", styles["Section"]))
    _append_pdf_table(
        story,
        [
            ["Indicador", "Valor"],
            ["Recursos HTML detectados", str(report["automaticAccessibilitySummary"]["htmlResourcesDetected"])],
            ["Recursos HTML analizados", str(report["automaticAccessibilitySummary"]["htmlResourcesAnalyzed"])],
            ["Recursos PDF detectados", str(report["automaticAccessibilitySummary"]["pdfResourcesDetected"])],
            ["Recursos PDF analizados", str(report["automaticAccessibilitySummary"]["pdfResourcesAnalyzed"])],
            ["Recursos Word detectados", str(report["automaticAccessibilitySummary"]["wordResourcesDetected"])],
            ["Recursos Word analizados", str(report["automaticAccessibilitySummary"]["wordResourcesAnalyzed"])],
            ["Recursos de vídeo detectados", str(report["automaticAccessibilitySummary"]["videoResourcesDetected"])],
            ["Recursos de vídeo analizados", str(report["automaticAccessibilitySummary"]["videoResourcesAnalyzed"])],
            ["Total PASS", str(report["automaticAccessibilitySummary"]["passCount"])],
            ["Total FAIL", str(report["automaticAccessibilitySummary"]["failCount"])],
            ["Total WARNING", str(report["automaticAccessibilitySummary"]["warningCount"])],
            ["Total NOT_APPLICABLE", str(report["automaticAccessibilitySummary"]["notApplicableCount"])],
            ["Total ERROR", str(report["automaticAccessibilitySummary"]["errorCount"])],
        ],
        widths=[9 * cm, 4.2 * cm],
    )

    story.append(_pdf_paragraph("Puntuación por módulo", styles["Section"]))
    if report["moduleScores"]:
        rows = [["Módulo/sección", "Score", "Prioridad", "Recursos", "Incidencias principales"]]
        for row in report["moduleScores"]:
            rows.append(
                [
                    row["moduleTitle"],
                    f"{row['score']}/100",
                    _priority_label(row["priority"]),
                    str(row["resourcesAnalyzed"]),
                    _clip_text("; ".join(row["mainIssues"]), 180),
                ]
            )
        _append_pdf_table(story, rows, widths=[4.4 * cm, 2 * cm, 2.2 * cm, 2 * cm, 5 * cm])
    else:
        story.append(_pdf_paragraph("No hay recursos analizados automáticamente para calcular score por módulo.", styles["Normal"]))

    story.append(_pdf_paragraph("Puntuación por recurso", styles["Section"]))
    if report["resourceScores"]:
        rows = [["Recurso", "Tipo", "Módulo", "Score", "Prioridad", "Incidencia principal"]]
        for row in report["resourceScores"]:
            rows.append(
                [
                    row["title"],
                    row["typeLabel"],
                    row["moduleTitle"],
                    f"{row['score']}/100",
                    _priority_label(row["priority"]),
                    _clip_text(row["mainIssue"], 140),
                ]
            )
        _append_pdf_table(story, rows, widths=[3.6 * cm, 1.5 * cm, 3.1 * cm, 1.7 * cm, 1.9 * cm, 3.8 * cm])
    else:
        story.append(_pdf_paragraph("No hay recursos analizados automáticamente para calcular score por recurso.", styles["Normal"]))

    story.append(_pdf_paragraph("Principales incidencias", styles["Section"]))
    if not report["issueSummary"]:
        story.append(_pdf_paragraph("No se han detectado incidencias FAIL o WARNING en los checks automáticos.", styles["Normal"]))
    else:
        rows = [["Tipo", "Check", "Estado", "Nº recursos", "Recomendación"]]
        for issue_group in report["issueSummary"]:
            rows.append(
                [
                    _auto_resource_type_label(issue_group["resourceType"]),
                    issue_group["checkTitle"],
                    issue_group["status"],
                    str(issue_group["resourceCount"]),
                    _clip_text(issue_group["recommendation"], 220),
                ]
            )
        _append_pdf_table(story, rows, widths=[1.8 * cm, 3.3 * cm, 2.2 * cm, 2 * cm, 6.2 * cm])

    story.extend([PageBreak(), _pdf_paragraph("Detalle técnico", styles["Section"])])
    story.append(_pdf_paragraph("Los checks PASS se omiten en esta sección para mantener el informe accionable.", styles["Normal"]))
    _append_pdf_technical_resources(story, styles, "Detalle por recurso HTML", report["htmlResources"])
    _append_pdf_technical_resources(story, styles, "Detalle por recurso PDF", report["pdfResources"])
    _append_pdf_technical_resources(story, styles, "Detalle por recurso Word", report["wordResources"])
    _append_pdf_technical_resources(
        story,
        styles,
        "Detalle por recurso de vídeo",
        report["videoResources"],
        intro=VIDEO_ANALYSIS_SCOPE_NOTE,
    )

    story.append(_pdf_paragraph("Recursos no analizables automáticamente", styles["Section"]))
    if not report["notAutomaticallyAnalyzable"]:
        story.append(
            _pdf_paragraph(
                "No hay recursos pendientes de análisis automático por autenticación, interacción o cobertura.",
                styles["Normal"],
            )
        )
    else:
        story.append(_pdf_paragraph(AUTO_NOT_ANALYZABLE_EXPLANATION, styles["Normal"]))
        rows = [["Recurso", "Tipo", "Motivo", "Módulo/sección"]]
        for resource in report["notAutomaticallyAnalyzable"]:
            rows.append(
                [
                    resource["title"],
                    resource["type"],
                    resource["reason"],
                    resource["moduleTitle"] or resource["coursePath"],
                ]
            )
        _append_pdf_table(story, rows, widths=[5 * cm, 2.2 * cm, 3.3 * cm, 5 * cm])

    story.append(_pdf_paragraph("Revisión manual complementaria", styles["Section"]))
    if not report["routes"]:
        story.append(_pdf_paragraph("No hay hallazgos FAIL o PENDING en la checklist manual.", styles["Normal"]))
    else:
        for route in report["routes"]:
            story.append(_pdf_paragraph(f"Ruta: {route['coursePath']}", styles["Heading2"]))
            for resource in route["resources"]:
                story.append(
                    _pdf_paragraph(
                        f"{resource['title']} | Tipo: {resource['type']} | Origen: {resource['origin']}",
                        styles["Normal"],
                    )
                )
                rows = [["Estado", "Severidad", "Descripción", "Cómo arreglarlo", "Notas"]]
                for issue in resource["fails"] + resource["pending"]:
                    rows.append(
                        [
                            issue["status"],
                            issue["severity"],
                            issue["description"],
                            issue["recommendation"] or "Sin recomendación disponible.",
                            issue.get("comment") or "-",
                        ]
                )
                _append_pdf_table(story, rows, widths=[1.8 * cm, 2.2 * cm, 4 * cm, 4.8 * cm, 2.7 * cm])

    story.append(_pdf_paragraph("Apéndice", styles["Section"]))
    for key, value in report["appendix"]["statusDefinitions"].items():
        story.append(_pdf_paragraph(f"- {key}: {value}", styles["Normal"]))
    story.append(_pdf_paragraph(f"Fecha: {_format_report_date(report['appendix']['createdAt'])}", styles["Normal"]))
    story.append(_pdf_paragraph(f"Versión del sistema: {report['appendix']['systemVersion']}", styles["Normal"]))

    document = SimpleDocTemplate(
        str(destination),
        pagesize=A4,
        leftMargin=1.6 * cm,
        rightMargin=1.6 * cm,
        topMargin=1.4 * cm,
        bottomMargin=1.4 * cm,
    )
    document.build(story)


def _convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False

    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "pdf:writer_pdf_Export",
                "--outdir",
                str(pdf_path.parent),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
    except (OSError, subprocess.CalledProcessError):
        return False

    generated_pdf = docx_path.with_suffix(".pdf")
    if generated_pdf.exists() and generated_pdf != pdf_path:
        generated_pdf.replace(pdf_path)
    return pdf_path.exists()


def _persist_files(settings: Settings, job_id: str, report: dict[str, Any]) -> tuple[Path, Path, Path]:
    json_path, docx_path, pdf_path = _canonical_paths(settings, job_id)
    json_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    _write_docx(docx_path, report, settings.report_brand_name)
    if not _convert_docx_to_pdf(docx_path, pdf_path):
        _write_pdf(pdf_path, report, settings.report_brand_name)
    return json_path, docx_path, pdf_path


def get_report_or_404(session: Session, job_id: str) -> ReportRecord:
    report = session.exec(select(ReportRecord).where(ReportRecord.job_id == job_id)).first()
    if report is None or report.payload is None:
        raise AppError(
            code="report_not_found",
            message="Todavia no existe un informe generado para este job.",
            status_code=status.HTTP_404_NOT_FOUND,
            job_id=job_id,
        )
    return report


def generate_job_report(
    session: Session,
    settings: Settings,
    job_id: str,
    *,
    include_pending: bool = True,
    only_fails: bool = False,
    canvas_client: Any | None = None,
    canvas_credentials: Any | None = None,
    course_id: str | None = None,
) -> dict[str, Any]:
    payload = _build_report_payload(
        session,
        settings,
        job_id,
        include_pending=include_pending,
        only_fails=only_fails,
        canvas_client=canvas_client,
        canvas_credentials=canvas_credentials,
        course_id=course_id,
    )
    _, docx_path, pdf_path = _persist_files(settings, job_id, payload)
    generated_at = datetime.fromisoformat(payload["createdAt"])
    record = session.exec(select(ReportRecord).where(ReportRecord.job_id == job_id)).first()

    if record is None:
        record = ReportRecord(job_id=job_id)

    record.resource_count = payload["stats"]["resources"]
    record.failed_item_count = payload["stats"]["fails"]
    record.generated_at = generated_at
    record.pdf_path = str(pdf_path)
    record.docx_path = str(docx_path)
    record.payload = payload
    session.add(record)
    session.commit()
    return payload


def load_job_report(session: Session, job_id: str) -> dict[str, Any]:
    return _normalize_report_payload(get_report_or_404(session, job_id).payload)


def _normalize_report_payload(payload: dict[str, Any]) -> dict[str, Any]:
    normalized = json.loads(json.dumps(payload))
    automatic_summary = normalized.get("automaticAccessibilitySummary")
    if isinstance(automatic_summary, dict):
        if "wordResourcesDetected" not in automatic_summary:
            automatic_summary["wordResourcesDetected"] = automatic_summary.pop("docxResourcesDetected", 0)
        else:
            automatic_summary.pop("docxResourcesDetected", None)
        if "wordResourcesAnalyzed" not in automatic_summary:
            automatic_summary["wordResourcesAnalyzed"] = automatic_summary.pop("docxResourcesAnalyzed", 0)
        else:
            automatic_summary.pop("docxResourcesAnalyzed", None)
        automatic_summary.setdefault("videoResourcesDetected", 0)
        automatic_summary.setdefault("videoResourcesAnalyzed", 0)

    if "wordAccessibilitySummary" not in normalized:
        normalized["wordAccessibilitySummary"] = normalized.pop("docxAccessibilitySummary", _empty_auto_summary())
    else:
        normalized.pop("docxAccessibilitySummary", None)
    if "wordResources" not in normalized:
        normalized["wordResources"] = normalized.pop("docxResources", [])
    else:
        normalized.pop("docxResources", None)
    normalized.setdefault("videoAccessibilitySummary", _empty_auto_summary())
    normalized.setdefault("videoResources", [])

    for collection_name in ("htmlResources", "pdfResources", "wordResources", "videoResources"):
        for resource in normalized.get(collection_name, []):
            if not isinstance(resource, dict):
                continue
            score_fields = _score_fields(resource.get("checks", []))
            resource.setdefault("score", score_fields["score"])
            resource.setdefault("priority", score_fields["priority"])
            resource.setdefault("mainIssue", score_fields["mainIssue"])

    normalized.setdefault(
        "resourceScores",
        _resource_score_rows(
            normalized.get("htmlResources", []),
            normalized.get("pdfResources", []),
            normalized.get("wordResources", []),
            normalized.get("videoResources", []),
        ),
    )
    normalized.setdefault("moduleScores", _module_score_rows(normalized["resourceScores"]))
    normalized.setdefault(
        "executiveSummary",
        _build_executive_summary(
            normalized.get("accessSummary", {}),
            normalized["resourceScores"],
            normalized.get("issueSummary", []),
            normalized.get("notAutomaticallyAnalyzable", []),
        ),
    )

    for collection_name in ("issueSummary", "keyIssues"):
        for item in normalized.get(collection_name, []):
            if isinstance(item, dict) and item.get("resourceType") == "DOCX":
                item["resourceType"] = "WORD"
    return normalized


def _empty_auto_summary() -> dict[str, int]:
    return {
        "resourcesDetected": 0,
        "resourcesAnalyzed": 0,
        "passCount": 0,
        "failCount": 0,
        "warningCount": 0,
        "notApplicableCount": 0,
        "errorCount": 0,
    }


def get_report_file_info(session: Session, settings: Settings, job_id: str, fmt: str) -> tuple[Path, str, str]:
    if fmt not in MEDIA_TYPES:
        raise AppError(
            code="invalid_format",
            message="Formato de descarga no soportado.",
            status_code=status.HTTP_404_NOT_FOUND,
            job_id=job_id,
        )

    record = get_report_or_404(session, job_id)
    json_path, canonical_docx, canonical_pdf = _canonical_paths(settings, job_id)

    if fmt == "json":
        file_path = json_path
    elif fmt == "docx":
        file_path = Path(record.docx_path) if record.docx_path else canonical_docx
    else:
        file_path = Path(record.pdf_path) if record.pdf_path else canonical_pdf

    if not file_path.exists():
        raise AppError(
            code="report_file_missing",
            message="El archivo solicitado del informe no esta disponible.",
            status_code=status.HTTP_404_NOT_FOUND,
            job_id=job_id,
        )

    download_name = _stable_filename(job_id, record.generated_at, fmt)
    return file_path, MEDIA_TYPES[fmt], download_name


def _legacy_groups_from_payload(session: Session, payload: dict[str, Any]) -> list[ReportGroup]:
    groups: list[ReportGroup] = []
    for resource in payload["resources"]:
        if not resource["fails"]:
            continue
        review_resource = session.get(Resource, resource["resourceId"])
        if review_resource is None:
            continue
        groups.append(
            ReportGroup(
                resource=_resource_response(review_resource),
                failures=[
                    ReportFailure(
                        itemId=item["itemKey"],
                        label=item["label"],
                        recommendation=item["recommendation"] or "Sin recomendacion disponible.",
                    )
                    for item in resource["fails"]
                ],
            )
        )
    return groups


def generate_report(session: Session, settings: Settings, job_id: str) -> GeneratedReportResponse:
    payload = generate_job_report(session, settings, job_id, include_pending=False, only_fails=True)
    return GeneratedReportResponse(
        jobId=job_id,
        resourceCount=payload["stats"]["resources"],
        failedItemCount=payload["stats"]["fails"],
        groups=_legacy_groups_from_payload(session, payload),
        generatedAt=datetime.fromisoformat(payload["createdAt"]),
        downloads=_legacy_downloads(job_id),
    )


def load_report(session: Session, job_id: str) -> GeneratedReportResponse:
    payload = load_job_report(session, job_id)
    return GeneratedReportResponse(
        jobId=job_id,
        resourceCount=payload["stats"]["resources"],
        failedItemCount=payload["stats"]["fails"],
        groups=_legacy_groups_from_payload(session, payload),
        generatedAt=datetime.fromisoformat(payload["createdAt"]),
        downloads=_legacy_downloads(job_id),
    )
